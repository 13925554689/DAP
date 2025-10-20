"""
多格式导出器 - Layer 4
多种格式导出 (Excel, PDF, Word, JSON)

核心功能：
1. 多格式数据导出
2. 模板化导出
3. 批量导出处理
4. 导出任务管理
5. 格式转换和优化
"""

import asyncio
import logging
import json
import sqlite3
import numpy as np
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple, Set, Union
from pathlib import Path
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor
import threading
from collections import defaultdict
import pickle
import hashlib
import uuid
import zipfile
import base64
import io
from enum import Enum

# Excel处理
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, NamedStyle
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.worksheet.table import Table, TableStyleInfo
    from openpyxl.utils.dataframe import dataframe_to_rows

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# PDF生成
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
        Image,
        PageBreak,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib import colors
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie
    from reportlab.graphics.charts.linecharts import HorizontalLineChart

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# 图像处理
try:
    import matplotlib.pyplot as plt
    import seaborn as sns

    plt.rcParams["font.sans-serif"] = ["SimHei"]
    plt.rcParams["axes.unicode_minus"] = False
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    import plotly.io as pio

    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

# 压缩
try:
    import zipfile
    import tarfile

    COMPRESSION_AVAILABLE = True
except ImportError:
    COMPRESSION_AVAILABLE = False


class ExportFormat(Enum):
    """导出格式"""

    EXCEL = "excel"
    PDF = "pdf"
    WORD = "word"
    JSON = "json"
    CSV = "csv"
    HTML = "html"
    XML = "xml"
    TXT = "txt"


class ExportType(Enum):
    """导出类型"""

    DATA_EXPORT = "data_export"  # 数据导出
    REPORT_EXPORT = "report_export"  # 报表导出
    ANALYSIS_EXPORT = "analysis_export"  # 分析导出
    BATCH_EXPORT = "batch_export"  # 批量导出


class CompressionType(Enum):
    """压缩类型"""

    NONE = "none"
    ZIP = "zip"
    TAR_GZ = "tar_gz"
    TAR_BZ2 = "tar_bz2"


@dataclass
class ExportTask:
    """导出任务"""

    task_id: str
    task_name: str
    export_type: ExportType
    export_format: ExportFormat
    data_source: str
    template_config: Dict[str, Any]
    output_config: Dict[str, Any]
    compression: CompressionType
    created_at: datetime
    started_at: Optional[datetime]
    completed_at: Optional[datetime]
    status: str
    progress: float
    file_path: str
    file_size: int
    error_message: str


@dataclass
class ExportTemplate:
    """导出模板"""

    template_id: str
    template_name: str
    export_format: ExportFormat
    template_config: Dict[str, Any]
    style_config: Dict[str, Any]
    created_at: datetime
    updated_at: datetime
    usage_count: int


@dataclass
class ExportResult:
    """导出结果"""

    task_id: str
    success: bool
    file_path: str
    file_size: int
    export_time: float
    records_exported: int
    format_info: Dict[str, Any]


class MultiFormatExporter:
    """多格式导出器"""

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.logger = logging.getLogger(__name__)

        # 数据库配置
        self.exporter_db_path = self.config.get(
            "exporter_db_path", "data/multi_format_exporter.db"
        )
        self.export_path = self.config.get("export_path", "exports/")
        self.temp_path = self.config.get("temp_path", "temp/")

        # 任务管理
        self.active_tasks = {}
        self.task_queue = asyncio.Queue()
        self.template_cache = {}

        # 样式配置
        self.default_styles = {
            "font_family": self.config.get("font_family", "Microsoft YaHei"),
            "font_size": self.config.get("font_size", 10),
            "header_color": self.config.get("header_color", "#4472C4"),
            "border_color": self.config.get("border_color", "#000000"),
            "alternate_row_color": self.config.get("alternate_row_color", "#F2F2F2"),
        }

        # 并发控制
        self.max_workers = self.config.get("max_workers", 4)
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)
        self._lock = threading.RLock()

        # 确保目录存在
        Path(self.export_path).mkdir(parents=True, exist_ok=True)
        Path(self.temp_path).mkdir(parents=True, exist_ok=True)

        # 初始化数据库
        self._init_database()

        # 加载导出模板
        asyncio.create_task(self._load_export_templates())

        # 启动任务处理器
        asyncio.create_task(self._task_processor())

    def _init_database(self):
        """初始化导出器数据库"""
        try:
            Path(self.exporter_db_path).parent.mkdir(parents=True, exist_ok=True)

            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()

                # 导出任务表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS export_tasks (
                        task_id TEXT PRIMARY KEY,
                        task_name TEXT NOT NULL,
                        export_type TEXT NOT NULL,
                        export_format TEXT NOT NULL,
                        data_source TEXT NOT NULL,
                        template_config TEXT,
                        output_config TEXT,
                        compression TEXT DEFAULT 'none',
                        created_at TEXT NOT NULL,
                        started_at TEXT,
                        completed_at TEXT,
                        status TEXT DEFAULT 'pending',
                        progress REAL DEFAULT 0,
                        file_path TEXT,
                        file_size INTEGER DEFAULT 0,
                        error_message TEXT
                    )
                """
                )

                # 导出模板表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS export_templates (
                        template_id TEXT PRIMARY KEY,
                        template_name TEXT NOT NULL,
                        export_format TEXT NOT NULL,
                        template_config TEXT NOT NULL,
                        style_config TEXT,
                        created_at TEXT NOT NULL,
                        updated_at TEXT NOT NULL,
                        usage_count INTEGER DEFAULT 0
                    )
                """
                )

                # 导出统计表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS export_statistics (
                        stat_id TEXT PRIMARY KEY,
                        export_format TEXT NOT NULL,
                        total_exports INTEGER DEFAULT 0,
                        total_size INTEGER DEFAULT 0,
                        avg_export_time REAL DEFAULT 0,
                        success_rate REAL DEFAULT 100,
                        last_updated TEXT DEFAULT CURRENT_TIMESTAMP
                    )
                """
                )

                # 用户导出偏好表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS user_export_preferences (
                        user_id TEXT PRIMARY KEY,
                        preferred_format TEXT,
                        default_template TEXT,
                        compression_preference TEXT,
                        style_preferences TEXT,
                        updated_at TEXT DEFAULT CURRENT_TIMESTAMP
                    )
                """
                )

                # 索引
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_tasks_status ON export_tasks (status)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_tasks_created ON export_tasks (created_at)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_templates_format ON export_templates (export_format)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_stats_format ON export_statistics (export_format)"
                )

                conn.commit()

            self.logger.info("多格式导出器数据库初始化完成")

        except Exception as e:
            self.logger.error(f"多格式导出器数据库初始化失败: {e}")
            raise

    async def _load_export_templates(self):
        """加载导出模板"""
        try:
            # 创建默认模板
            await self._create_default_templates()

            # 从数据库加载模板
            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT template_id, template_name, export_format,
                           template_config, style_config, usage_count
                    FROM export_templates
                """
                )

                templates_loaded = 0
                for row in cursor.fetchall():
                    template_id = row[0]
                    try:
                        template = ExportTemplate(
                            template_id=template_id,
                            template_name=row[1],
                            export_format=ExportFormat(row[2]),
                            template_config=json.loads(row[3]),
                            style_config=json.loads(row[4]) if row[4] else {},
                            created_at=datetime.now(),
                            updated_at=datetime.now(),
                            usage_count=row[5],
                        )

                        with self._lock:
                            self.template_cache[template_id] = template

                        templates_loaded += 1

                    except Exception as e:
                        self.logger.error(f"加载模板 {template_id} 失败: {e}")

            self.logger.info(f"导出模板加载完成: {templates_loaded} 个模板")

        except Exception as e:
            self.logger.error(f"加载导出模板失败: {e}")

    async def _create_default_templates(self):
        """创建默认导出模板"""
        try:
            # Excel模板
            excel_template = {
                "template_id": "excel_standard_template",
                "template_name": "Excel标准模板",
                "export_format": "excel",
                "template_config": {
                    "include_header": True,
                    "include_index": False,
                    "freeze_panes": True,
                    "auto_filter": True,
                    "sheet_name": "数据导出",
                    "table_style": "TableStyleMedium2",
                },
                "style_config": {
                    "header_font": {"bold": True, "color": "FFFFFF"},
                    "header_fill": {"color": "4472C4"},
                    "data_font": {"size": 10},
                    "border": {"style": "thin", "color": "000000"},
                },
            }

            # PDF模板
            pdf_template = {
                "template_id": "pdf_report_template",
                "template_name": "PDF报表模板",
                "export_format": "pdf",
                "template_config": {
                    "page_size": "A4",
                    "orientation": "portrait",
                    "margins": {"top": 2, "bottom": 2, "left": 2, "right": 2},
                    "include_header": True,
                    "include_footer": True,
                    "font_name": "Helvetica",
                },
                "style_config": {
                    "title_style": {"fontSize": 16, "textColor": "black"},
                    "header_style": {
                        "fontSize": 12,
                        "textColor": "white",
                        "backColor": "blue",
                    },
                    "data_style": {"fontSize": 10, "textColor": "black"},
                },
            }

            # Word模板
            word_template = {
                "template_id": "word_document_template",
                "template_name": "Word文档模板",
                "export_format": "word",
                "template_config": {
                    "include_toc": False,
                    "page_orientation": "portrait",
                    "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
                    "table_style": "Table Grid",
                },
                "style_config": {
                    "title_font": {"size": 16, "bold": True},
                    "header_font": {"size": 12, "bold": True},
                    "body_font": {"size": 10},
                },
            }

            templates = [excel_template, pdf_template, word_template]

            for template_data in templates:
                await self._save_template_to_db(template_data)

        except Exception as e:
            self.logger.error(f"创建默认模板失败: {e}")

    async def _save_template_to_db(self, template_data: Dict[str, Any]):
        """保存模板到数据库"""
        try:
            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    INSERT OR REPLACE INTO export_templates
                    (template_id, template_name, export_format, template_config,
                     style_config, created_at, updated_at, usage_count)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """,
                    (
                        template_data["template_id"],
                        template_data["template_name"],
                        template_data["export_format"],
                        json.dumps(template_data["template_config"]),
                        json.dumps(template_data.get("style_config", {})),
                        datetime.now().isoformat(),
                        datetime.now().isoformat(),
                        0,
                    ),
                )
                conn.commit()

        except Exception as e:
            self.logger.error(f"保存模板到数据库失败: {e}")

    async def create_export_task(self, export_config: Dict[str, Any]) -> Dict[str, Any]:
        """创建导出任务"""
        try:
            task_id = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"

            task = ExportTask(
                task_id=task_id,
                task_name=export_config.get("task_name", "数据导出"),
                export_type=ExportType(export_config.get("export_type", "data_export")),
                export_format=ExportFormat(export_config["export_format"]),
                data_source=export_config["data_source"],
                template_config=export_config.get("template_config", {}),
                output_config=export_config.get("output_config", {}),
                compression=CompressionType(export_config.get("compression", "none")),
                created_at=datetime.now(),
                started_at=None,
                completed_at=None,
                status="pending",
                progress=0.0,
                file_path="",
                file_size=0,
                error_message="",
            )

            # 保存到数据库
            await self._save_task_to_db(task)

            # 添加到任务队列
            await self.task_queue.put(task)

            # 更新活跃任务
            with self._lock:
                self.active_tasks[task_id] = task

            result = {
                "status": "success",
                "task_id": task_id,
                "message": f"导出任务 '{task.task_name}' 创建成功",
            }

            self.logger.info(f"导出任务创建成功: {task_id}")
            return result

        except Exception as e:
            self.logger.error(f"创建导出任务失败: {e}")
            return {"status": "error", "error": str(e)}

    async def _save_task_to_db(self, task: ExportTask):
        """保存任务到数据库"""
        try:
            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    INSERT OR REPLACE INTO export_tasks
                    (task_id, task_name, export_type, export_format, data_source,
                     template_config, output_config, compression, created_at,
                     started_at, completed_at, status, progress, file_path,
                     file_size, error_message)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                    (
                        task.task_id,
                        task.task_name,
                        task.export_type.value,
                        task.export_format.value,
                        task.data_source,
                        json.dumps(task.template_config),
                        json.dumps(task.output_config),
                        task.compression.value,
                        task.created_at.isoformat(),
                        task.started_at.isoformat() if task.started_at else None,
                        task.completed_at.isoformat() if task.completed_at else None,
                        task.status,
                        task.progress,
                        task.file_path,
                        task.file_size,
                        task.error_message,
                    ),
                )
                conn.commit()

        except Exception as e:
            self.logger.error(f"保存任务到数据库失败: {e}")

    async def _task_processor(self):
        """任务处理器"""
        try:
            while True:
                try:
                    # 获取任务
                    task = await self.task_queue.get()

                    # 处理任务
                    await self._process_export_task(task)

                    # 标记任务完成
                    self.task_queue.task_done()

                except Exception as e:
                    self.logger.error(f"任务处理失败: {e}")
                    await asyncio.sleep(1)

        except Exception as e:
            self.logger.error(f"任务处理器失败: {e}")

    async def _process_export_task(self, task: ExportTask):
        """处理导出任务"""
        try:
            # 更新任务状态
            task.status = "running"
            task.started_at = datetime.now()
            task.progress = 0.0

            await self._update_task_progress(task)

            # 获取数据
            data = await self._get_export_data(task.data_source)
            if data is None:
                task.status = "failed"
                task.error_message = "无法获取数据源"
                await self._save_task_to_db(task)
                return

            task.progress = 20.0
            await self._update_task_progress(task)

            # 根据格式导出
            export_result = None

            if task.export_format == ExportFormat.EXCEL:
                export_result = await self._export_to_excel(data, task)
            elif task.export_format == ExportFormat.PDF:
                export_result = await self._export_to_pdf(data, task)
            elif task.export_format == ExportFormat.WORD:
                export_result = await self._export_to_word(data, task)
            elif task.export_format == ExportFormat.JSON:
                export_result = await self._export_to_json(data, task)
            elif task.export_format == ExportFormat.CSV:
                export_result = await self._export_to_csv(data, task)
            elif task.export_format == ExportFormat.HTML:
                export_result = await self._export_to_html(data, task)
            else:
                task.status = "failed"
                task.error_message = f"不支持的导出格式: {task.export_format.value}"
                await self._save_task_to_db(task)
                return

            task.progress = 80.0
            await self._update_task_progress(task)

            if export_result and export_result.success:
                # 压缩文件（如果需要）
                if task.compression != CompressionType.NONE:
                    compressed_path = await self._compress_file(
                        export_result.file_path, task.compression
                    )
                    if compressed_path:
                        export_result.file_path = compressed_path
                        export_result.file_size = Path(compressed_path).stat().st_size

                task.status = "completed"
                task.file_path = export_result.file_path
                task.file_size = export_result.file_size
                task.progress = 100.0
            else:
                task.status = "failed"
                task.error_message = "导出失败"

            task.completed_at = datetime.now()
            await self._save_task_to_db(task)

            # 更新统计
            await self._update_export_statistics(task)

            self.logger.info(f"导出任务完成: {task.task_id}, 状态: {task.status}")

        except Exception as e:
            self.logger.error(f"处理导出任务失败: {e}")
            task.status = "failed"
            task.error_message = str(e)
            task.completed_at = datetime.now()
            await self._save_task_to_db(task)

    async def _get_export_data(self, data_source: str) -> Optional[Any]:
        """获取导出数据"""
        try:
            # 这里应该根据数据源获取实际数据
            # 为了演示，生成示例数据

            if data_source == "financial_data":
                # 财务数据示例
                data = {
                    "balance_sheet": pd.DataFrame(
                        {
                            "科目代码": ["1001", "1002", "1122", "1403", "1601"],
                            "科目名称": ["库存现金", "银行存款", "应收账款", "原材料", "固定资产"],
                            "期初余额": [50000, 800000, 200000, 150000, 2000000],
                            "期末余额": [60000, 950000, 180000, 180000, 1950000],
                            "变动金额": [10000, 150000, -20000, 30000, -50000],
                        }
                    ),
                    "income_statement": pd.DataFrame(
                        {
                            "科目代码": ["4001", "5001", "5201", "5202"],
                            "科目名称": ["主营业务收入", "主营业务成本", "销售费用", "管理费用"],
                            "本期金额": [5000000, 3000000, 500000, 300000],
                            "上期金额": [4500000, 2700000, 450000, 280000],
                            "变动比例": ["11.1%", "11.1%", "11.1%", "7.1%"],
                        }
                    ),
                }

            elif data_source == "audit_results":
                # 审计结果示例
                data = pd.DataFrame(
                    {
                        "审计事项": ["现金盘点", "应收账款函证", "存货监盘", "收入截止测试", "费用完整性测试"],
                        "审计结果": ["无异常", "发现差异", "无异常", "无异常", "发现遗漏"],
                        "风险等级": ["低", "中", "低", "低", "高"],
                        "建议措施": ["继续保持", "加强内控", "定期盘点", "规范流程", "完善制度"],
                        "跟进状态": ["已完成", "整改中", "已完成", "已完成", "待整改"],
                    }
                )

            else:
                # 默认示例数据
                data = pd.DataFrame(
                    {
                        "ID": range(1, 101),
                        "名称": [f"项目{i}" for i in range(1, 101)],
                        "数值": np.random.randint(1000, 10000, 100),
                        "类别": np.random.choice(["A", "B", "C"], 100),
                        "日期": pd.date_range("2023-01-01", periods=100, freq="D"),
                    }
                )

            return data

        except Exception as e:
            self.logger.error(f"获取导出数据失败: {e}")
            return None

    async def _export_to_excel(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到Excel"""
        try:
            if not OPENPYXL_AVAILABLE:
                raise ImportError("openpyxl 不可用，无法导出Excel")

            start_time = datetime.now()

            # 生成文件名
            filename = (
                f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            )
            file_path = Path(self.export_path) / filename

            # 创建工作簿
            wb = openpyxl.Workbook()

            if isinstance(data, dict):
                # 多表数据
                # 删除默认工作表
                wb.remove(wb.active)

                for sheet_name, sheet_data in data.items():
                    ws = wb.create_sheet(title=sheet_name)
                    await self._add_dataframe_to_worksheet(ws, sheet_data, task)

            elif isinstance(data, pd.DataFrame):
                # 单表数据
                ws = wb.active
                ws.title = task.output_config.get("sheet_name", "数据导出")
                await self._add_dataframe_to_worksheet(ws, data, task)

            else:
                raise ValueError("不支持的数据格式")

            # 保存文件
            wb.save(file_path)

            execution_time = (datetime.now() - start_time).total_seconds()
            file_size = file_path.stat().st_size
            records_count = (
                len(data)
                if isinstance(data, pd.DataFrame)
                else sum(
                    len(df) for df in data.values() if isinstance(df, pd.DataFrame)
                )
            )

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                export_time=execution_time,
                records_exported=records_count,
                format_info={"format": "xlsx", "sheets": wb.sheetnames},
            )

        except Exception as e:
            self.logger.error(f"Excel导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _add_dataframe_to_worksheet(
        self, ws, data: pd.DataFrame, task: ExportTask
    ):
        """将DataFrame添加到工作表"""
        try:
            if data.empty:
                return

            # 获取模板配置
            template_config = task.template_config
            include_header = template_config.get("include_header", True)
            include_index = template_config.get("include_index", False)

            # 写入数据
            for r in dataframe_to_rows(
                data, index=include_index, header=include_header
            ):
                ws.append(r)

            # 应用样式
            if include_header:
                header_row = 1
                for col in range(
                    1, len(data.columns) + (1 if include_index else 0) + 1
                ):
                    cell = ws.cell(row=header_row, column=col)
                    cell.font = Font(bold=True, color="FFFFFF")
                    cell.fill = PatternFill(
                        start_color="4472C4", end_color="4472C4", fill_type="solid"
                    )
                    cell.alignment = Alignment(horizontal="center")

            # 自动调整列宽
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column_letter].width = adjusted_width

            # 添加筛选器
            if template_config.get("auto_filter", True) and not data.empty:
                ws.auto_filter.ref = ws.dimensions

            # 冻结窗格
            if template_config.get("freeze_panes", True) and include_header:
                ws.freeze_panes = "A2"

        except Exception as e:
            self.logger.error(f"添加DataFrame到工作表失败: {e}")

    async def _export_to_pdf(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到PDF"""
        try:
            if not REPORTLAB_AVAILABLE:
                raise ImportError("reportlab 不可用，无法导出PDF")

            start_time = datetime.now()

            # 生成文件名
            filename = (
                f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )
            file_path = Path(self.export_path) / filename

            # 创建PDF文档
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            story = []

            # 样式
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle", parent=styles["Heading1"], alignment=1
            )

            # 添加标题
            story.append(Paragraph(task.task_name, title_style))
            story.append(Spacer(1, 12))

            # 添加生成时间
            story.append(
                Paragraph(
                    f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            if isinstance(data, dict):
                # 多表数据
                for table_name, table_data in data.items():
                    if isinstance(table_data, pd.DataFrame) and not table_data.empty:
                        story.append(Paragraph(table_name, styles["Heading2"]))
                        story.append(Spacer(1, 6))

                        # 创建表格
                        table = await self._create_pdf_table(table_data)
                        if table:
                            story.append(table)
                            story.append(Spacer(1, 12))

            elif isinstance(data, pd.DataFrame):
                # 单表数据
                table = await self._create_pdf_table(data)
                if table:
                    story.append(table)

            # 生成PDF
            doc.build(story)

            execution_time = (datetime.now() - start_time).total_seconds()
            file_size = file_path.stat().st_size
            records_count = (
                len(data)
                if isinstance(data, pd.DataFrame)
                else sum(
                    len(df) for df in data.values() if isinstance(df, pd.DataFrame)
                )
            )

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                export_time=execution_time,
                records_exported=records_count,
                format_info={"format": "pdf", "pages": 1},
            )

        except Exception as e:
            self.logger.error(f"PDF导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _create_pdf_table(self, data: pd.DataFrame) -> Optional[Table]:
        """创建PDF表格"""
        try:
            if data.empty:
                return None

            # 准备表格数据
            table_data = [data.columns.tolist()]

            # 限制行数以避免PDF过大
            max_rows = 100
            for _, row in data.head(max_rows).iterrows():
                table_data.append([str(cell) for cell in row])

            # 创建表格
            table = Table(table_data)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("FONTSIZE", (0, 0), (-1, 0), 10),
                        ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                        ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ]
                )
            )

            return table

        except Exception as e:
            self.logger.error(f"创建PDF表格失败: {e}")
            return None

    async def _export_to_word(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到Word"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                raise ImportError("python-docx 不可用，无法导出Word")

            start_time = datetime.now()

            # 生成文件名
            filename = (
                f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            file_path = Path(self.export_path) / filename

            # 创建文档
            doc = Document()

            # 添加标题
            title = doc.add_heading(task.task_name, 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            if isinstance(data, dict):
                # 多表数据
                for table_name, table_data in data.items():
                    if isinstance(table_data, pd.DataFrame) and not table_data.empty:
                        doc.add_heading(table_name, level=1)
                        await self._add_dataframe_to_word_table(doc, table_data)

            elif isinstance(data, pd.DataFrame):
                # 单表数据
                await self._add_dataframe_to_word_table(doc, data)

            # 保存文档
            doc.save(file_path)

            execution_time = (datetime.now() - start_time).total_seconds()
            file_size = file_path.stat().st_size
            records_count = (
                len(data)
                if isinstance(data, pd.DataFrame)
                else sum(
                    len(df) for df in data.values() if isinstance(df, pd.DataFrame)
                )
            )

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                export_time=execution_time,
                records_exported=records_count,
                format_info={"format": "docx", "paragraphs": len(doc.paragraphs)},
            )

        except Exception as e:
            self.logger.error(f"Word导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _add_dataframe_to_word_table(self, doc, data: pd.DataFrame):
        """将DataFrame添加到Word表格"""
        try:
            if data.empty:
                return

            # 限制行数
            max_rows = 100
            display_data = data.head(max_rows)

            # 创建表格
            table = doc.add_table(rows=1, cols=len(display_data.columns))
            table.style = "Table Grid"

            # 添加表头
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(display_data.columns):
                hdr_cells[i].text = str(col_name)

            # 添加数据行
            for _, row in display_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

        except Exception as e:
            self.logger.error(f"添加DataFrame到Word表格失败: {e}")

    async def _export_to_json(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到JSON"""
        try:
            start_time = datetime.now()

            # 生成文件名
            filename = (
                f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            )
            file_path = Path(self.export_path) / filename

            # 准备JSON数据
            if isinstance(data, dict):
                json_data = {}
                for key, value in data.items():
                    if isinstance(value, pd.DataFrame):
                        json_data[key] = value.to_dict("records")
                    else:
                        json_data[key] = value
            elif isinstance(data, pd.DataFrame):
                json_data = data.to_dict("records")
            else:
                json_data = data

            # 写入JSON文件
            with open(file_path, "w", encoding="utf-8") as f:
                json.dump(json_data, f, ensure_ascii=False, indent=2, default=str)

            execution_time = (datetime.now() - start_time).total_seconds()
            file_size = file_path.stat().st_size
            records_count = (
                len(data)
                if isinstance(data, pd.DataFrame)
                else sum(
                    len(df) for df in data.values() if isinstance(df, pd.DataFrame)
                )
            )

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                export_time=execution_time,
                records_exported=records_count,
                format_info={"format": "json", "encoding": "utf-8"},
            )

        except Exception as e:
            self.logger.error(f"JSON导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _export_to_csv(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到CSV"""
        try:
            start_time = datetime.now()

            if isinstance(data, dict):
                # 多文件CSV
                file_paths = []
                total_records = 0

                for table_name, table_data in data.items():
                    if isinstance(table_data, pd.DataFrame):
                        filename = f"{task.task_name}_{table_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        file_path = Path(self.export_path) / filename
                        table_data.to_csv(file_path, index=False, encoding="utf-8-sig")
                        file_paths.append(str(file_path))
                        total_records += len(table_data)

                # 如果有多个文件，打包成ZIP
                if len(file_paths) > 1:
                    zip_filename = f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
                    zip_path = Path(self.export_path) / zip_filename

                    with zipfile.ZipFile(zip_path, "w") as zipf:
                        for file_path in file_paths:
                            zipf.write(file_path, Path(file_path).name)
                            Path(file_path).unlink()  # 删除原始文件

                    final_path = str(zip_path)
                    file_size = zip_path.stat().st_size
                else:
                    final_path = file_paths[0] if file_paths else ""
                    file_size = Path(final_path).stat().st_size if final_path else 0

            elif isinstance(data, pd.DataFrame):
                # 单文件CSV
                filename = (
                    f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                )
                final_path = str(Path(self.export_path) / filename)
                data.to_csv(final_path, index=False, encoding="utf-8-sig")
                file_size = Path(final_path).stat().st_size
                total_records = len(data)

            else:
                raise ValueError("不支持的数据格式")

            execution_time = (datetime.now() - start_time).total_seconds()

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=final_path,
                file_size=file_size,
                export_time=execution_time,
                records_exported=total_records,
                format_info={"format": "csv", "encoding": "utf-8-sig"},
            )

        except Exception as e:
            self.logger.error(f"CSV导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _export_to_html(self, data: Any, task: ExportTask) -> ExportResult:
        """导出到HTML"""
        try:
            start_time = datetime.now()

            # 生成文件名
            filename = (
                f"{task.task_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            )
            file_path = Path(self.export_path) / filename

            # 创建HTML内容
            html_content = f"""
            <!DOCTYPE html>
            <html lang="zh-CN">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{task.task_name}</title>
                <style>
                    body {{ font-family: 'Microsoft YaHei', Arial, sans-serif; margin: 20px; }}
                    .header {{ text-align: center; margin-bottom: 30px; }}
                    .title {{ font-size: 24px; font-weight: bold; color: #333; }}
                    .timestamp {{ color: #666; margin-top: 10px; }}
                    .table-container {{ margin: 20px 0; }}
                    .table-title {{ font-size: 18px; font-weight: bold; margin-bottom: 10px; color: #4472C4; }}
                    table {{ width: 100%; border-collapse: collapse; margin-bottom: 30px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #4472C4; color: white; font-weight: bold; }}
                    tr:nth-child(even) {{ background-color: #f2f2f2; }}
                    tr:hover {{ background-color: #e8f4f8; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <div class="title">{task.task_name}</div>
                    <div class="timestamp">生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
                </div>
            """

            records_count = 0

            if isinstance(data, dict):
                # 多表数据
                for table_name, table_data in data.items():
                    if isinstance(table_data, pd.DataFrame) and not table_data.empty:
                        html_content += f"""
                        <div class="table-container">
                            <div class="table-title">{table_name}</div>
                            {table_data.to_html(classes='data-table', table_id=f'table-{table_name}', escape=False)}
                        </div>
                        """
                        records_count += len(table_data)

            elif isinstance(data, pd.DataFrame):
                # 单表数据
                html_content += f"""
                <div class="table-container">
                    {data.to_html(classes='data-table', table_id='main-table', escape=False)}
                </div>
                """
                records_count = len(data)

            html_content += """
            </body>
            </html>
            """

            # 写入HTML文件
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_content)

            execution_time = (datetime.now() - start_time).total_seconds()
            file_size = file_path.stat().st_size

            return ExportResult(
                task_id=task.task_id,
                success=True,
                file_path=str(file_path),
                file_size=file_size,
                export_time=execution_time,
                records_exported=records_count,
                format_info={"format": "html", "encoding": "utf-8"},
            )

        except Exception as e:
            self.logger.error(f"HTML导出失败: {e}")
            return ExportResult(
                task_id=task.task_id,
                success=False,
                file_path="",
                file_size=0,
                export_time=0,
                records_exported=0,
                format_info={"error": str(e)},
            )

    async def _compress_file(
        self, file_path: str, compression_type: CompressionType
    ) -> Optional[str]:
        """压缩文件"""
        try:
            if compression_type == CompressionType.NONE:
                return file_path

            source_path = Path(file_path)

            if compression_type == CompressionType.ZIP:
                compressed_path = source_path.with_suffix(".zip")
                with zipfile.ZipFile(
                    compressed_path, "w", zipfile.ZIP_DEFLATED
                ) as zipf:
                    zipf.write(source_path, source_path.name)

            elif compression_type == CompressionType.TAR_GZ:
                compressed_path = source_path.with_suffix(".tar.gz")
                with tarfile.open(compressed_path, "w:gz") as tar:
                    tar.add(source_path, source_path.name)

            elif compression_type == CompressionType.TAR_BZ2:
                compressed_path = source_path.with_suffix(".tar.bz2")
                with tarfile.open(compressed_path, "w:bz2") as tar:
                    tar.add(source_path, source_path.name)

            else:
                return file_path

            # 删除原始文件
            source_path.unlink()

            return str(compressed_path)

        except Exception as e:
            self.logger.error(f"文件压缩失败: {e}")
            return file_path

    async def _update_task_progress(self, task: ExportTask):
        """更新任务进度"""
        try:
            with self._lock:
                if task.task_id in self.active_tasks:
                    self.active_tasks[task.task_id] = task

            # 可以在这里添加进度通知逻辑

        except Exception as e:
            self.logger.error(f"更新任务进度失败: {e}")

    async def _update_export_statistics(self, task: ExportTask):
        """更新导出统计"""
        try:
            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()

                # 获取现有统计
                cursor.execute(
                    """
                    SELECT total_exports, total_size, avg_export_time, success_rate
                    FROM export_statistics
                    WHERE export_format = ?
                """,
                    (task.export_format.value,),
                )

                result = cursor.fetchone()
                if result:
                    total_exports, total_size, avg_export_time, success_rate = result

                    # 更新统计
                    new_total_exports = total_exports + 1
                    new_total_size = total_size + task.file_size

                    execution_time = (
                        (task.completed_at - task.started_at).total_seconds()
                        if task.started_at and task.completed_at
                        else 0
                    )
                    new_avg_export_time = (
                        avg_export_time * total_exports + execution_time
                    ) / new_total_exports

                    success_count = int(success_rate * total_exports / 100) + (
                        1 if task.status == "completed" else 0
                    )
                    new_success_rate = success_count / new_total_exports * 100

                    cursor.execute(
                        """
                        UPDATE export_statistics
                        SET total_exports = ?, total_size = ?, avg_export_time = ?,
                            success_rate = ?, last_updated = ?
                        WHERE export_format = ?
                    """,
                        (
                            new_total_exports,
                            new_total_size,
                            new_avg_export_time,
                            new_success_rate,
                            datetime.now().isoformat(),
                            task.export_format.value,
                        ),
                    )
                else:
                    # 插入新统计
                    execution_time = (
                        (task.completed_at - task.started_at).total_seconds()
                        if task.started_at and task.completed_at
                        else 0
                    )
                    success_rate = 100.0 if task.status == "completed" else 0.0

                    cursor.execute(
                        """
                        INSERT INTO export_statistics
                        (stat_id, export_format, total_exports, total_size,
                         avg_export_time, success_rate, last_updated)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """,
                        (
                            f"stat_{task.export_format.value}",
                            task.export_format.value,
                            1,
                            task.file_size,
                            execution_time,
                            success_rate,
                            datetime.now().isoformat(),
                        ),
                    )

                conn.commit()

        except Exception as e:
            self.logger.error(f"更新导出统计失败: {e}")

    async def get_task_status(self, task_id: str) -> Dict[str, Any]:
        """获取任务状态"""
        try:
            with self._lock:
                if task_id in self.active_tasks:
                    task = self.active_tasks[task_id]
                    return {
                        "task_id": task_id,
                        "status": task.status,
                        "progress": task.progress,
                        "file_path": task.file_path,
                        "file_size": task.file_size,
                        "error_message": task.error_message,
                        "created_at": task.created_at.isoformat(),
                        "started_at": task.started_at.isoformat()
                        if task.started_at
                        else None,
                        "completed_at": task.completed_at.isoformat()
                        if task.completed_at
                        else None,
                    }

            # 从数据库查询
            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT task_name, status, progress, file_path, file_size,
                           error_message, created_at, started_at, completed_at
                    FROM export_tasks
                    WHERE task_id = ?
                """,
                    (task_id,),
                )

                result = cursor.fetchone()
                if result:
                    return {
                        "task_id": task_id,
                        "task_name": result[0],
                        "status": result[1],
                        "progress": result[2],
                        "file_path": result[3],
                        "file_size": result[4],
                        "error_message": result[5],
                        "created_at": result[6],
                        "started_at": result[7],
                        "completed_at": result[8],
                    }

            return {"status": "error", "error": f"任务 {task_id} 不存在"}

        except Exception as e:
            self.logger.error(f"获取任务状态失败: {e}")
            return {"status": "error", "error": str(e)}

    async def list_export_tasks(
        self, filter_config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """列出导出任务"""
        try:
            filter_config = filter_config or {}

            with sqlite3.connect(self.exporter_db_path) as conn:
                cursor = conn.cursor()

                # 构建查询条件
                where_clause = "WHERE 1=1"
                params = []

                if filter_config.get("status"):
                    where_clause += " AND status = ?"
                    params.append(filter_config["status"])

                if filter_config.get("export_format"):
                    where_clause += " AND export_format = ?"
                    params.append(filter_config["export_format"])

                # 执行查询
                cursor.execute(
                    f"""
                    SELECT task_id, task_name, export_format, status, progress,
                           file_size, created_at, completed_at
                    FROM export_tasks
                    {where_clause}
                    ORDER BY created_at DESC
                    LIMIT ?
                """,
                    params + [filter_config.get("limit", 50)],
                )

                tasks = []
                for row in cursor.fetchall():
                    tasks.append(
                        {
                            "task_id": row[0],
                            "task_name": row[1],
                            "export_format": row[2],
                            "status": row[3],
                            "progress": row[4],
                            "file_size": row[5],
                            "created_at": row[6],
                            "completed_at": row[7],
                        }
                    )

                return {"status": "success", "tasks": tasks, "total_count": len(tasks)}

        except Exception as e:
            self.logger.error(f"列出导出任务失败: {e}")
            return {"status": "error", "error": str(e)}

    async def __aenter__(self):
        """异步上下文管理器入口"""
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """异步上下文管理器出口"""
        await self.cleanup()

    async def cleanup(self):
        """清理资源"""
        try:
            if hasattr(self, "executor"):
                self.executor.shutdown(wait=True)

            self.logger.info("多格式导出器资源清理完成")

        except Exception as e:
            self.logger.error(f"资源清理失败: {e}")


async def main():
    """测试主函数"""
    config = {
        "exporter_db_path": "data/test_multi_format_exporter.db",
        "export_path": "test_exports/",
        "temp_path": "test_temp/",
    }

    async with MultiFormatExporter(config) as exporter:
        # 创建Excel导出任务
        excel_config = {
            "task_name": "财务数据导出",
            "export_type": "data_export",
            "export_format": "excel",
            "data_source": "financial_data",
            "template_config": {
                "include_header": True,
                "auto_filter": True,
                "freeze_panes": True,
            },
            "compression": "none",
        }

        task_result = await exporter.create_export_task(excel_config)
        print(f"导出任务创建结果: {json.dumps(task_result, indent=2, ensure_ascii=False)}")

        if task_result["status"] == "success":
            task_id = task_result["task_id"]

            # 等待任务完成
            while True:
                await asyncio.sleep(2)
                status = await exporter.get_task_status(task_id)
                print(f"任务状态: {status['status']}, 进度: {status['progress']:.1f}%")

                if status["status"] in ["completed", "failed"]:
                    break

            print(f"最终状态: {json.dumps(status, indent=2, ensure_ascii=False)}")

        # 列出所有任务
        tasks_list = await exporter.list_export_tasks()
        print(f"导出任务列表: {json.dumps(tasks_list, indent=2, ensure_ascii=False)}")


if __name__ == "__main__":
    asyncio.run(main())

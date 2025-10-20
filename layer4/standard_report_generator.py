"""
标准报表生成器 - Layer 4
基于审计模板的报表生成

核心功能：
1. 多种审计报表模板
2. 动态数据填充和计算
3. 多格式输出（Excel、PDF、Word、HTML）
4. 智能图表生成
5. 报表质量检查和验证
"""

import asyncio
import logging
import json
import sqlite3
import numpy as np
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple, Set, Union
from pathlib import Path
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor
import threading
from collections import defaultdict, Counter
import pickle
import hashlib
import base64
import io
from enum import Enum

from utils.async_utils import schedule_async_task

# Excel处理
try:
    import openpyxl
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.chart import BarChart, LineChart, PieChart, Reference
    from openpyxl.utils.dataframe import dataframe_to_rows

    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# PDF生成
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import (
        SimpleDocTemplate,
        Table,
        TableStyle,
        Paragraph,
        Spacer,
        Image,
    )
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.graphics.shapes import Drawing
    from reportlab.graphics.charts.barcharts import VerticalBarChart
    from reportlab.graphics.charts.piecharts import Pie

    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.shared import OxmlElement, qn

    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# 图表生成
try:
    import matplotlib.pyplot as plt
    import seaborn as sns

    plt.rcParams["font.sans-serif"] = ["SimHei"]  # 中文字体
    plt.rcParams["axes.unicode_minus"] = False
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    from plotly.subplots import make_subplots
    import plotly.io as pio

    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False


class ReportType(Enum):
    """报表类型"""

    FINANCIAL_STATEMENT = "financial_statement"  # 财务报表
    AUDIT_SUMMARY = "audit_summary"  # 审计汇总
    EXCEPTION_REPORT = "exception_report"  # 异常报告
    COMPLIANCE_REPORT = "compliance_report"  # 合规报告
    MANAGEMENT_LETTER = "management_letter"  # 管理建议书
    RISK_ASSESSMENT = "risk_assessment"  # 风险评估报告


class OutputFormat(Enum):
    """输出格式"""

    EXCEL = "excel"
    PDF = "pdf"
    WORD = "word"
    HTML = "html"
    JSON = "json"


@dataclass
class ReportTemplate:
    """报表模板"""

    template_id: str
    template_name: str
    report_type: ReportType
    template_structure: Dict[str, Any]
    data_requirements: List[str]
    chart_definitions: List[Dict[str, Any]]
    style_settings: Dict[str, Any]
    created_at: datetime
    updated_at: datetime
    version: str


@dataclass
class ReportSection:
    """报表章节"""

    section_id: str
    section_name: str
    section_type: str  # table, chart, text, summary
    data_source: str
    template_content: Dict[str, Any]
    order_index: int


@dataclass
class GeneratedReport:
    """生成的报表"""

    report_id: str
    template_id: str
    report_name: str
    output_format: OutputFormat
    file_path: str
    generated_at: datetime
    data_period: str
    metadata: Dict[str, Any]
    file_size: int


class StandardReportGenerator:
    """标准报表生成器"""

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.logger = logging.getLogger(__name__)

        # 数据库配置
        self.reports_db_path = self.config.get("reports_db_path", "data/reports.db")
        self.templates_path = self.config.get("templates_path", "data/templates/")
        self.output_path = self.config.get("output_path", "exports/")

        # 模板缓存
        self.template_cache = {}
        self.chart_cache = {}

        # 样式配置
        self.default_styles = {
            "font_family": self.config.get("font_family", "SimHei"),
            "font_size": self.config.get("font_size", 10),
            "header_color": self.config.get("header_color", "#4472C4"),
            "border_color": self.config.get("border_color", "#000000"),
            "alternate_row_color": self.config.get("alternate_row_color", "#F2F2F2"),
        }

        # 并发控制
        self.max_workers = self.config.get("max_workers", 4)
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)
        self._lock = threading.RLock()

        # 确保目录存在
        Path(self.templates_path).mkdir(parents=True, exist_ok=True)
        Path(self.output_path).mkdir(parents=True, exist_ok=True)

        # 初始化数据库
        self._init_database()

        # 加载报表模板
        schedule_async_task(
            self._load_report_templates,
            logger=self.logger,
            task_name="load_report_templates",
        )

    def _init_database(self):
        """初始化报表数据库"""
        try:
            Path(self.reports_db_path).parent.mkdir(parents=True, exist_ok=True)

            with sqlite3.connect(self.reports_db_path) as conn:
                cursor = conn.cursor()

                # 报表模板表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS report_templates (
                        template_id TEXT PRIMARY KEY,
                        template_name TEXT NOT NULL,
                        report_type TEXT NOT NULL,
                        template_structure TEXT NOT NULL,
                        data_requirements TEXT,
                        chart_definitions TEXT,
                        style_settings TEXT,
                        created_at TEXT NOT NULL,
                        updated_at TEXT NOT NULL,
                        version TEXT DEFAULT '1.0',
                        is_active BOOLEAN DEFAULT 1
                    )
                """
                )

                # 报表章节表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS report_sections (
                        section_id TEXT PRIMARY KEY,
                        template_id TEXT NOT NULL,
                        section_name TEXT NOT NULL,
                        section_type TEXT NOT NULL,
                        data_source TEXT,
                        template_content TEXT,
                        order_index INTEGER DEFAULT 0,
                        FOREIGN KEY (template_id) REFERENCES report_templates (template_id)
                    )
                """
                )

                # 生成报表记录表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS generated_reports (
                        report_id TEXT PRIMARY KEY,
                        template_id TEXT NOT NULL,
                        report_name TEXT NOT NULL,
                        output_format TEXT NOT NULL,
                        file_path TEXT NOT NULL,
                        generated_at TEXT NOT NULL,
                        data_period TEXT,
                        metadata TEXT,
                        file_size INTEGER DEFAULT 0,
                        FOREIGN KEY (template_id) REFERENCES report_templates (template_id)
                    )
                """
                )

                # 报表使用统计表
                cursor.execute(
                    """
                    CREATE TABLE IF NOT EXISTS report_usage_stats (
                        stat_id TEXT PRIMARY KEY,
                        template_id TEXT NOT NULL,
                        usage_count INTEGER DEFAULT 0,
                        last_generated TEXT,
                        avg_generation_time REAL DEFAULT 0,
                        success_rate REAL DEFAULT 100,
                        user_ratings REAL DEFAULT 0,
                        FOREIGN KEY (template_id) REFERENCES report_templates (template_id)
                    )
                """
                )

                # 索引
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_templates_type ON report_templates (report_type)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_sections_template ON report_sections (template_id)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_reports_template ON generated_reports (template_id)"
                )
                cursor.execute(
                    "CREATE INDEX IF NOT EXISTS idx_reports_generated ON generated_reports (generated_at)"
                )

                conn.commit()

            self.logger.info("报表数据库初始化完成")

        except Exception as e:
            self.logger.error(f"报表数据库初始化失败: {e}")
            raise

    async def _load_report_templates(self):
        """加载报表模板"""
        try:
            # 创建默认模板
            await self._create_default_templates()

            # 从数据库加载现有模板
            with sqlite3.connect(self.reports_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT template_id, template_name, report_type, template_structure,
                           data_requirements, chart_definitions, style_settings
                    FROM report_templates
                    WHERE is_active = 1
                """
                )

                templates_loaded = 0
                for row in cursor.fetchall():
                    template_id = row[0]
                    try:
                        template = ReportTemplate(
                            template_id=template_id,
                            template_name=row[1],
                            report_type=ReportType(row[2]),
                            template_structure=json.loads(row[3]),
                            data_requirements=json.loads(row[4]) if row[4] else [],
                            chart_definitions=json.loads(row[5]) if row[5] else [],
                            style_settings=json.loads(row[6]) if row[6] else {},
                            created_at=datetime.now(),
                            updated_at=datetime.now(),
                            version="1.0",
                        )

                        with self._lock:
                            self.template_cache[template_id] = template

                        templates_loaded += 1

                    except Exception as e:
                        self.logger.error(f"加载模板 {template_id} 失败: {e}")

            self.logger.info(f"报表模板加载完成: {templates_loaded} 个模板")

        except Exception as e:
            self.logger.error(f"加载报表模板失败: {e}")

    async def _create_default_templates(self):
        """创建默认报表模板"""
        try:
            # 财务报表模板
            financial_template = {
                "template_id": "financial_statement_template",
                "template_name": "财务报表汇总",
                "report_type": "financial_statement",
                "template_structure": {
                    "title": "财务报表分析汇总",
                    "sections": [
                        {
                            "name": "资产负债表分析",
                            "type": "table",
                            "data_source": "balance_sheet",
                            "columns": ["科目名称", "本期金额", "上期金额", "变动金额", "变动比例"],
                        },
                        {
                            "name": "利润表分析",
                            "type": "table",
                            "data_source": "income_statement",
                            "columns": ["科目名称", "本期金额", "上期金额", "变动金额", "变动比例"],
                        },
                        {
                            "name": "财务指标分析",
                            "type": "chart",
                            "chart_type": "bar",
                            "data_source": "financial_ratios",
                        },
                    ],
                },
                "data_requirements": [
                    "balance_sheet",
                    "income_statement",
                    "financial_ratios",
                ],
                "chart_definitions": [
                    {
                        "chart_id": "assets_trend",
                        "chart_type": "line",
                        "title": "资产趋势分析",
                        "x_axis": "月份",
                        "y_axis": "金额",
                    }
                ],
            }

            # 异常报告模板
            exception_template = {
                "template_id": "exception_report_template",
                "template_name": "异常事项报告",
                "report_type": "exception_report",
                "template_structure": {
                    "title": "数据异常检测报告",
                    "sections": [
                        {
                            "name": "异常汇总",
                            "type": "summary",
                            "data_source": "anomaly_summary",
                        },
                        {
                            "name": "异常明细",
                            "type": "table",
                            "data_source": "anomaly_details",
                            "columns": ["异常类型", "检测时间", "异常描述", "严重程度", "影响范围"],
                        },
                        {
                            "name": "异常分布图",
                            "type": "chart",
                            "chart_type": "pie",
                            "data_source": "anomaly_distribution",
                        },
                    ],
                },
                "data_requirements": [
                    "anomaly_summary",
                    "anomaly_details",
                    "anomaly_distribution",
                ],
                "chart_definitions": [
                    {
                        "chart_id": "anomaly_trend",
                        "chart_type": "line",
                        "title": "异常趋势分析",
                        "x_axis": "时间",
                        "y_axis": "异常数量",
                    }
                ],
            }

            # 审计汇总模板
            audit_summary_template = {
                "template_id": "audit_summary_template",
                "template_name": "审计工作汇总",
                "report_type": "audit_summary",
                "template_structure": {
                    "title": "审计工作执行情况汇总",
                    "sections": [
                        {"name": "审计范围", "type": "text", "data_source": "audit_scope"},
                        {
                            "name": "执行的审计程序",
                            "type": "table",
                            "data_source": "audit_procedures",
                            "columns": ["程序名称", "执行状态", "发现问题数", "完成时间"],
                        },
                        {
                            "name": "发现问题统计",
                            "type": "chart",
                            "chart_type": "bar",
                            "data_source": "issues_statistics",
                        },
                    ],
                },
                "data_requirements": [
                    "audit_scope",
                    "audit_procedures",
                    "issues_statistics",
                ],
                "chart_definitions": [],
            }

            templates = [financial_template, exception_template, audit_summary_template]

            for template_data in templates:
                await self._save_template_to_db(template_data)

        except Exception as e:
            self.logger.error(f"创建默认模板失败: {e}")

    async def _save_template_to_db(self, template_data: Dict[str, Any]):
        """保存模板到数据库"""
        try:
            with sqlite3.connect(self.reports_db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    INSERT OR REPLACE INTO report_templates
                    (template_id, template_name, report_type, template_structure,
                     data_requirements, chart_definitions, style_settings,
                     created_at, updated_at, version)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                    (
                        template_data["template_id"],
                        template_data["template_name"],
                        template_data["report_type"],
                        json.dumps(template_data["template_structure"]),
                        json.dumps(template_data["data_requirements"]),
                        json.dumps(template_data.get("chart_definitions", [])),
                        json.dumps(template_data.get("style_settings", {})),
                        datetime.now().isoformat(),
                        datetime.now().isoformat(),
                        "1.0",
                    ),
                )
                conn.commit()

        except Exception as e:
            self.logger.error(f"保存模板到数据库失败: {e}")

    async def generate_report(
        self, generation_config: Dict[str, Any]
    ) -> Dict[str, Any]:
        """生成报表"""
        try:
            template_id = generation_config["template_id"]
            output_format = OutputFormat(
                generation_config.get("output_format", "excel")
            )
            data = generation_config["data"]

            generation_id = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{hashlib.md5(str(generation_config).encode()).hexdigest()[:8]}"

            result = {
                "generation_id": generation_id,
                "template_id": template_id,
                "started_at": datetime.now().isoformat(),
                "output_format": output_format.value,
                "file_path": "",
                "file_size": 0,
                "sections_generated": 0,
                "charts_generated": 0,
                "validation_results": {},
                "status": "success",
            }

            # 获取模板
            template = self.template_cache.get(template_id)
            if not template:
                return {
                    "status": "error",
                    "error": f"模板 {template_id} 不存在",
                    "started_at": datetime.now().isoformat(),
                }

            # 验证数据完整性
            validation_result = await self._validate_report_data(template, data)
            result["validation_results"] = validation_result

            if not validation_result["is_valid"]:
                return {
                    "status": "error",
                    "error": "数据验证失败",
                    "validation_results": validation_result,
                    "started_at": datetime.now().isoformat(),
                }

            generation_start_time = datetime.now()

            # 根据输出格式生成报表
            if output_format == OutputFormat.EXCEL:
                file_path = await self._generate_excel_report(
                    template, data, generation_config
                )
            elif output_format == OutputFormat.PDF:
                file_path = await self._generate_pdf_report(
                    template, data, generation_config
                )
            elif output_format == OutputFormat.WORD:
                file_path = await self._generate_word_report(
                    template, data, generation_config
                )
            elif output_format == OutputFormat.HTML:
                file_path = await self._generate_html_report(
                    template, data, generation_config
                )
            else:
                return {
                    "status": "error",
                    "error": f"不支持的输出格式: {output_format.value}",
                    "started_at": datetime.now().isoformat(),
                }

            # 计算文件大小
            if file_path and Path(file_path).exists():
                file_size = Path(file_path).stat().st_size
            else:
                file_size = 0

            result["file_path"] = file_path
            result["file_size"] = file_size
            result["sections_generated"] = len(
                template.template_structure.get("sections", [])
            )
            result["charts_generated"] = len(template.chart_definitions)

            generation_time = (datetime.now() - generation_start_time).total_seconds()
            result["generation_time"] = generation_time
            result["completed_at"] = datetime.now().isoformat()

            # 保存生成记录
            await self._save_generation_record(result, template, generation_config)

            # 更新使用统计
            await self._update_template_usage_stats(template_id, generation_time, True)

            self.logger.info(f"报表生成完成: {file_path}, 大小: {file_size} 字节")

            return result

        except Exception as e:
            self.logger.error(f"报表生成失败: {e}")
            return {
                "status": "error",
                "error": str(e),
                "started_at": datetime.now().isoformat(),
            }

    async def _validate_report_data(
        self, template: ReportTemplate, data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """验证报表数据"""
        try:
            validation_result = {
                "is_valid": True,
                "missing_data": [],
                "data_quality_issues": [],
                "warnings": [],
            }

            # 检查必需数据源
            for required_source in template.data_requirements:
                if required_source not in data or not data[required_source]:
                    validation_result["missing_data"].append(required_source)
                    validation_result["is_valid"] = False

            # 检查数据质量
            for source_name, source_data in data.items():
                if isinstance(source_data, list):
                    if len(source_data) == 0:
                        validation_result["warnings"].append(f"{source_name} 数据为空")

                    # 检查数据一致性
                    for item in source_data[:5]:  # 检查前5条
                        if isinstance(item, dict):
                            for key, value in item.items():
                                if value is None:
                                    validation_result["data_quality_issues"].append(
                                        f"{source_name}.{key} 包含空值"
                                    )

            return validation_result

        except Exception as e:
            self.logger.error(f"数据验证失败: {e}")
            return {"is_valid": False, "error": str(e)}

    async def _generate_excel_report(
        self, template: ReportTemplate, data: Dict[str, Any], config: Dict[str, Any]
    ) -> str:
        """生成Excel报表"""
        try:
            if not OPENPYXL_AVAILABLE:
                raise ImportError("openpyxl 不可用，无法生成Excel报表")

            # 创建工作簿
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = template.template_name

            current_row = 1

            # 添加标题
            title = template.template_structure.get("title", template.template_name)
            ws.merge_cells(f"A{current_row}:F{current_row}")
            title_cell = ws[f"A{current_row}"]
            title_cell.value = title
            title_cell.font = Font(size=16, bold=True)
            title_cell.alignment = Alignment(horizontal="center")
            current_row += 2

            # 添加生成时间
            ws[
                f"A{current_row}"
            ] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            current_row += 2

            # 处理各个章节
            sections = template.template_structure.get("sections", [])
            for section in sections:
                current_row = await self._add_excel_section(
                    ws, section, data, current_row
                )
                current_row += 2  # 章节间隔

            # 添加图表
            for chart_def in template.chart_definitions:
                current_row = await self._add_excel_chart(
                    ws, chart_def, data, current_row
                )
                current_row += 15  # 图表高度

            # 保存文件
            filename = f"{template.template_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            file_path = Path(self.output_path) / filename
            wb.save(file_path)

            return str(file_path)

        except Exception as e:
            self.logger.error(f"生成Excel报表失败: {e}")
            raise

    async def _add_excel_section(
        self, ws, section: Dict[str, Any], data: Dict[str, Any], start_row: int
    ) -> int:
        """添加Excel章节"""
        try:
            current_row = start_row

            # 添加章节标题
            ws[f"A{current_row}"] = section["name"]
            ws[f"A{current_row}"].font = Font(size=12, bold=True)
            current_row += 1

            section_type = section["type"]
            data_source = section.get("data_source")

            if section_type == "table" and data_source in data:
                current_row = await self._add_excel_table(
                    ws, section, data[data_source], current_row
                )
            elif section_type == "summary":
                current_row = await self._add_excel_summary(
                    ws, section, data.get(data_source, {}), current_row
                )
            elif section_type == "text":
                current_row = await self._add_excel_text(
                    ws, section, data.get(data_source, ""), current_row
                )

            return current_row

        except Exception as e:
            self.logger.error(f"添加Excel章节失败: {e}")
            return start_row

    async def _add_excel_table(
        self, ws, section: Dict[str, Any], table_data: Any, start_row: int
    ) -> int:
        """添加Excel表格"""
        try:
            current_row = start_row

            if isinstance(table_data, list) and table_data:
                # 如果指定了列名，使用指定的列名
                columns = section.get("columns")
                if columns:
                    # 写入表头
                    for col_idx, col_name in enumerate(columns, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=col_name)
                        cell.font = Font(bold=True)
                        cell.fill = PatternFill(
                            start_color=self.default_styles["header_color"][1:],
                            end_color=self.default_styles["header_color"][1:],
                            fill_type="solid",
                        )

                    current_row += 1

                    # 写入数据
                    for row_data in table_data:
                        if isinstance(row_data, dict):
                            for col_idx, col_name in enumerate(columns, 1):
                                value = row_data.get(col_name, "")
                                ws.cell(row=current_row, column=col_idx, value=value)
                        current_row += 1

                elif isinstance(table_data[0], dict):
                    # 使用数据的键作为列名
                    columns = list(table_data[0].keys())

                    # 写入表头
                    for col_idx, col_name in enumerate(columns, 1):
                        cell = ws.cell(row=current_row, column=col_idx, value=col_name)
                        cell.font = Font(bold=True)

                    current_row += 1

                    # 写入数据
                    for row_data in table_data:
                        for col_idx, col_name in enumerate(columns, 1):
                            value = row_data.get(col_name, "")
                            ws.cell(row=current_row, column=col_idx, value=value)
                        current_row += 1

            return current_row

        except Exception as e:
            self.logger.error(f"添加Excel表格失败: {e}")
            return start_row

    async def _add_excel_summary(
        self, ws, section: Dict[str, Any], summary_data: Dict[str, Any], start_row: int
    ) -> int:
        """添加Excel汇总"""
        try:
            current_row = start_row

            for key, value in summary_data.items():
                ws[f"A{current_row}"] = key
                ws[f"B{current_row}"] = value
                current_row += 1

            return current_row

        except Exception as e:
            self.logger.error(f"添加Excel汇总失败: {e}")
            return start_row

    async def _add_excel_text(
        self, ws, section: Dict[str, Any], text_data: str, start_row: int
    ) -> int:
        """添加Excel文本"""
        try:
            ws[f"A{start_row}"] = text_data
            return start_row + 1

        except Exception as e:
            self.logger.error(f"添加Excel文本失败: {e}")
            return start_row

    async def _add_excel_chart(
        self, ws, chart_def: Dict[str, Any], data: Dict[str, Any], start_row: int
    ) -> int:
        """添加Excel图表"""
        try:
            if not OPENPYXL_AVAILABLE:
                return start_row

            chart_type = chart_def.get("chart_type", "bar")
            title = chart_def.get("title", "图表")
            data_source = chart_def.get("data_source")

            if not data_source or data_source not in data:
                return start_row

            chart_data = data[data_source]

            # 根据图表类型创建图表
            if chart_type == "bar":
                chart = BarChart()
            elif chart_type == "line":
                chart = LineChart()
            elif chart_type == "pie":
                chart = PieChart()
            else:
                return start_row

            chart.title = title

            # 这里简化处理，实际应该根据数据结构动态生成
            # 添加图表到工作表
            ws.add_chart(chart, f"A{start_row}")

            return start_row

        except Exception as e:
            self.logger.error(f"添加Excel图表失败: {e}")
            return start_row

    async def _generate_pdf_report(
        self, template: ReportTemplate, data: Dict[str, Any], config: Dict[str, Any]
    ) -> str:
        """生成PDF报表"""
        try:
            if not REPORTLAB_AVAILABLE:
                raise ImportError("reportlab 不可用，无法生成PDF报表")

            filename = f"{template.template_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = Path(self.output_path) / filename

            # 创建PDF文档
            doc = SimpleDocTemplate(str(file_path), pagesize=A4)
            story = []

            # 样式
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                "CustomTitle", parent=styles["Heading1"], alignment=1
            )

            # 添加标题
            title = template.template_structure.get("title", template.template_name)
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))

            # 添加生成时间
            story.append(
                Paragraph(
                    f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            # 处理各个章节
            sections = template.template_structure.get("sections", [])
            for section in sections:
                await self._add_pdf_section(story, section, data, styles)

            # 生成PDF
            doc.build(story)

            return str(file_path)

        except Exception as e:
            self.logger.error(f"生成PDF报表失败: {e}")
            raise

    async def _add_pdf_section(
        self, story: List, section: Dict[str, Any], data: Dict[str, Any], styles
    ):
        """添加PDF章节"""
        try:
            # 添加章节标题
            story.append(Paragraph(section["name"], styles["Heading2"]))
            story.append(Spacer(1, 6))

            section_type = section["type"]
            data_source = section.get("data_source")

            if section_type == "table" and data_source in data:
                await self._add_pdf_table(story, section, data[data_source])
            elif section_type == "summary":
                await self._add_pdf_summary(
                    story, section, data.get(data_source, {}), styles
                )
            elif section_type == "text":
                await self._add_pdf_text(
                    story, section, data.get(data_source, ""), styles
                )

            story.append(Spacer(1, 12))

        except Exception as e:
            self.logger.error(f"添加PDF章节失败: {e}")

    async def _add_pdf_table(
        self, story: List, section: Dict[str, Any], table_data: Any
    ):
        """添加PDF表格"""
        try:
            if isinstance(table_data, list) and table_data:
                # 准备表格数据
                table_rows = []

                columns = section.get("columns")
                if columns:
                    table_rows.append(columns)

                    for row_data in table_data[:20]:  # 限制行数
                        if isinstance(row_data, dict):
                            row = [str(row_data.get(col, "")) for col in columns]
                            table_rows.append(row)

                # 创建表格
                if table_rows:
                    table = Table(table_rows)
                    table.setStyle(
                        TableStyle(
                            [
                                ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                                ("FONTSIZE", (0, 0), (-1, 0), 10),
                                ("BOTTOMPADDING", (0, 0), (-1, 0), 12),
                                ("BACKGROUND", (0, 1), (-1, -1), colors.beige),
                                ("GRID", (0, 0), (-1, -1), 1, colors.black),
                            ]
                        )
                    )

                    story.append(table)

        except Exception as e:
            self.logger.error(f"添加PDF表格失败: {e}")

    async def _add_pdf_summary(
        self, story: List, section: Dict[str, Any], summary_data: Dict[str, Any], styles
    ):
        """添加PDF汇总"""
        try:
            for key, value in summary_data.items():
                story.append(Paragraph(f"{key}: {value}", styles["Normal"]))

        except Exception as e:
            self.logger.error(f"添加PDF汇总失败: {e}")

    async def _add_pdf_text(
        self, story: List, section: Dict[str, Any], text_data: str, styles
    ):
        """添加PDF文本"""
        try:
            story.append(Paragraph(text_data, styles["Normal"]))

        except Exception as e:
            self.logger.error(f"添加PDF文本失败: {e}")

    async def _generate_word_report(
        self, template: ReportTemplate, data: Dict[str, Any], config: Dict[str, Any]
    ) -> str:
        """生成Word报表"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                raise ImportError("python-docx 不可用，无法生成Word报表")

            # 创建文档
            doc = Document()

            # 添加标题
            title = template.template_structure.get("title", template.template_name)
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加生成时间
            doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # 处理各个章节
            sections = template.template_structure.get("sections", [])
            for section in sections:
                await self._add_word_section(doc, section, data)

            # 保存文件
            filename = f"{template.template_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            file_path = Path(self.output_path) / filename
            doc.save(file_path)

            return str(file_path)

        except Exception as e:
            self.logger.error(f"生成Word报表失败: {e}")
            raise

    async def _add_word_section(
        self, doc, section: Dict[str, Any], data: Dict[str, Any]
    ):
        """添加Word章节"""
        try:
            # 添加章节标题
            doc.add_heading(section["name"], level=1)

            section_type = section["type"]
            data_source = section.get("data_source")

            if section_type == "table" and data_source in data:
                await self._add_word_table(doc, section, data[data_source])
            elif section_type == "summary":
                await self._add_word_summary(doc, section, data.get(data_source, {}))
            elif section_type == "text":
                await self._add_word_text(doc, section, data.get(data_source, ""))

        except Exception as e:
            self.logger.error(f"添加Word章节失败: {e}")

    async def _add_word_table(self, doc, section: Dict[str, Any], table_data: Any):
        """添加Word表格"""
        try:
            if isinstance(table_data, list) and table_data:
                columns = section.get("columns")
                if columns and table_data:
                    # 创建表格
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = "Table Grid"

                    # 添加表头
                    hdr_cells = table.rows[0].cells
                    for i, col_name in enumerate(columns):
                        hdr_cells[i].text = col_name

                    # 添加数据行
                    for row_data in table_data[:50]:  # 限制行数
                        if isinstance(row_data, dict):
                            row_cells = table.add_row().cells
                            for i, col_name in enumerate(columns):
                                row_cells[i].text = str(row_data.get(col_name, ""))

        except Exception as e:
            self.logger.error(f"添加Word表格失败: {e}")

    async def _add_word_summary(
        self, doc, section: Dict[str, Any], summary_data: Dict[str, Any]
    ):
        """添加Word汇总"""
        try:
            for key, value in summary_data.items():
                doc.add_paragraph(f"{key}: {value}")

        except Exception as e:
            self.logger.error(f"添加Word汇总失败: {e}")

    async def _add_word_text(self, doc, section: Dict[str, Any], text_data: str):
        """添加Word文本"""
        try:
            doc.add_paragraph(text_data)

        except Exception as e:
            self.logger.error(f"添加Word文本失败: {e}")

    async def _generate_html_report(
        self, template: ReportTemplate, data: Dict[str, Any], config: Dict[str, Any]
    ) -> str:
        """生成HTML报表"""
        try:
            # HTML模板
            html_content = f"""
            <!DOCTYPE html>
            <html lang="zh-CN">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>{template.template_name}</title>
                <style>
                    body {{ font-family: 'SimHei', sans-serif; margin: 20px; }}
                    .header {{ text-align: center; margin-bottom: 30px; }}
                    .title {{ font-size: 24px; font-weight: bold; margin-bottom: 10px; }}
                    .timestamp {{ color: #666; margin-bottom: 20px; }}
                    .section {{ margin-bottom: 30px; }}
                    .section-title {{ font-size: 18px; font-weight: bold; margin-bottom: 15px; background-color: #f0f0f0; padding: 10px; }}
                    table {{ width: 100%; border-collapse: collapse; margin-bottom: 20px; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #4472C4; color: white; }}
                    tr:nth-child(even) {{ background-color: #f2f2f2; }}
                    .summary-item {{ margin: 5px 0; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <div class="title">{template.template_structure.get('title', template.template_name)}</div>
                    <div class="timestamp">生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
                </div>
            """

            # 处理各个章节
            sections = template.template_structure.get("sections", [])
            for section in sections:
                html_content += await self._generate_html_section(section, data)

            html_content += """
            </body>
            </html>
            """

            # 保存文件
            filename = f"{template.template_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            file_path = Path(self.output_path) / filename

            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_content)

            return str(file_path)

        except Exception as e:
            self.logger.error(f"生成HTML报表失败: {e}")
            raise

    async def _generate_html_section(
        self, section: Dict[str, Any], data: Dict[str, Any]
    ) -> str:
        """生成HTML章节"""
        try:
            html = f'<div class="section"><div class="section-title">{section["name"]}</div>'

            section_type = section["type"]
            data_source = section.get("data_source")

            if section_type == "table" and data_source in data:
                html += await self._generate_html_table(section, data[data_source])
            elif section_type == "summary":
                html += await self._generate_html_summary(
                    section, data.get(data_source, {})
                )
            elif section_type == "text":
                html += f'<p>{data.get(data_source, "")}</p>'

            html += "</div>"
            return html

        except Exception as e:
            self.logger.error(f"生成HTML章节失败: {e}")
            return ""

    async def _generate_html_table(
        self, section: Dict[str, Any], table_data: Any
    ) -> str:
        """生成HTML表格"""
        try:
            if not isinstance(table_data, list) or not table_data:
                return "<p>无数据</p>"

            html = "<table>"

            columns = section.get("columns")
            if columns:
                # 添加表头
                html += "<tr>"
                for col in columns:
                    html += f"<th>{col}</th>"
                html += "</tr>"

                # 添加数据行
                for row_data in table_data[:100]:  # 限制行数
                    if isinstance(row_data, dict):
                        html += "<tr>"
                        for col in columns:
                            value = row_data.get(col, "")
                            html += f"<td>{value}</td>"
                        html += "</tr>"

            html += "</table>"
            return html

        except Exception as e:
            self.logger.error(f"生成HTML表格失败: {e}")
            return ""

    async def _generate_html_summary(
        self, section: Dict[str, Any], summary_data: Dict[str, Any]
    ) -> str:
        """生成HTML汇总"""
        try:
            html = ""
            for key, value in summary_data.items():
                html += (
                    f'<div class="summary-item"><strong>{key}:</strong> {value}</div>'
                )
            return html

        except Exception as e:
            self.logger.error(f"生成HTML汇总失败: {e}")
            return ""

    async def list_templates(
        self, filter_config: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """列出报表模板"""
        try:
            filter_config = filter_config or {}

            result = {"templates": [], "total_count": 0, "filtered_count": 0}

            # 获取所有模板
            templates = []
            for template_id, template in self.template_cache.items():
                template_info = {
                    "template_id": template_id,
                    "template_name": template.template_name,
                    "report_type": template.report_type.value,
                    "data_requirements": template.data_requirements,
                    "sections_count": len(
                        template.template_structure.get("sections", [])
                    ),
                    "charts_count": len(template.chart_definitions),
                    "version": template.version,
                }

                # 应用过滤器
                if (
                    filter_config.get("report_type")
                    and template.report_type.value != filter_config["report_type"]
                ):
                    continue

                templates.append(template_info)

            result["templates"] = templates
            result["total_count"] = len(self.template_cache)
            result["filtered_count"] = len(templates)

            return result

        except Exception as e:
            self.logger.error(f"列出模板失败: {e}")
            return {"status": "error", "error": str(e)}

    async def __aenter__(self):
        """异步上下文管理器入口"""
        return self

    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """异步上下文管理器出口"""
        await self.cleanup()

    async def cleanup(self):
        """清理资源"""
        try:
            if hasattr(self, "executor"):
                self.executor.shutdown(wait=True)

            self.logger.info("标准报表生成器资源清理完成")

        except Exception as e:
            self.logger.error(f"资源清理失败: {e}")


async def main():
    """测试主函数"""
    config = {
        "reports_db_path": "data/test_reports.db",
        "templates_path": "data/test_templates/",
        "output_path": "test_exports/",
    }

    async with StandardReportGenerator(config) as generator:
        # 测试数据
        test_data = {
            "balance_sheet": [
                {
                    "科目名称": "货币资金",
                    "本期金额": 1000000,
                    "上期金额": 800000,
                    "变动金额": 200000,
                    "变动比例": "25%",
                },
                {
                    "科目名称": "应收账款",
                    "本期金额": 500000,
                    "上期金额": 600000,
                    "变动金额": -100000,
                    "变动比例": "-16.7%",
                },
                {
                    "科目名称": "存货",
                    "本期金额": 300000,
                    "上期金额": 250000,
                    "变动金额": 50000,
                    "变动比例": "20%",
                },
            ],
            "income_statement": [
                {
                    "科目名称": "主营业务收入",
                    "本期金额": 5000000,
                    "上期金额": 4500000,
                    "变动金额": 500000,
                    "变动比例": "11.1%",
                },
                {
                    "科目名称": "主营业务成本",
                    "本期金额": 3000000,
                    "上期金额": 2800000,
                    "变动金额": 200000,
                    "变动比例": "7.1%",
                },
            ],
            "financial_ratios": {
                "流动比率": 2.5,
                "速动比率": 1.8,
                "资产负债率": 0.45,
                "净资产收益率": 0.15,
            },
        }

        # 生成Excel报表
        generation_config = {
            "template_id": "financial_statement_template",
            "output_format": "excel",
            "data": test_data,
            "report_name": "财务报表分析",
        }

        result = await generator.generate_report(generation_config)
        print(f"报表生成结果: {json.dumps(result, indent=2, ensure_ascii=False)}")

        # 列出模板
        templates = await generator.list_templates()
        print(f"可用模板: {json.dumps(templates, indent=2, ensure_ascii=False)}")


if __name__ == "__main__":
    asyncio.run(main())

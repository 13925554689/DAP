"""
增强的导出服务 - 审计证据专业导出
支持PDF、Word、图谱可视化等多种格式

核心功能:
1. 专业审计PDF报告导出(支持中文、表格、图表)
2. Word审计底稿导出(符合审计规范)
3. 证据关系图谱可视化导出
4. 批量导出管理
5. 导出模板管理
"""

import asyncio
import logging
import json
import sqlite3
import numpy as np
import pandas as pd
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple, Set, Union
from pathlib import Path
from dataclasses import dataclass, asdict
from concurrent.futures import ThreadPoolExecutor
import threading
from collections import defaultdict
import uuid
import zipfile
import io

# PDF生成 (中文支持)
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm, mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, Paragraph,
        Spacer, Image, PageBreak, KeepTogether, Frame, PageTemplate
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# Word文档生成
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

# 图谱可视化
try:
    import networkx as nx
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')  # 无GUI后端
    plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False
    NETWORKX_AVAILABLE = True
except ImportError:
    NETWORKX_AVAILABLE = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False


class ExportFormat:
    """导出格式"""
    PDF = "pdf"
    WORD = "word"
    EXCEL = "excel"
    GRAPH = "graph"
    HTML = "html"
    JSON = "json"


class AuditReportTemplate:
    """审计报告模板"""
    STANDARD = "standard"  # 标准审计报告
    SUMMARY = "summary"  # 摘要报告
    DETAILED = "detailed"  # 详细报告
    EVIDENCE = "evidence"  # 证据汇总
    FINDINGS = "findings"  # 发现汇总


@dataclass
class ExportTask:
    """导出任务"""
    task_id: str
    task_name: str
    export_format: str
    template_type: str
    data_source: Dict[str, Any]
    output_path: str
    status: str = "pending"
    progress: float = 0.0
    created_at: datetime = None
    completed_at: Optional[datetime] = None
    file_path: str = ""
    file_size: int = 0
    error_message: str = ""


class EnhancedExportService:
    """增强的导出服务"""

    def __init__(self, config: Dict[str, Any] = None):
        self.config = config or {}
        self.logger = logging.getLogger(__name__)

        # 路径配置
        self.export_path = Path(self.config.get("export_path", "exports/"))
        self.template_path = Path(self.config.get("template_path", "templates/"))
        self.temp_path = Path(self.config.get("temp_path", "temp/"))

        # 确保目录存在
        self.export_path.mkdir(parents=True, exist_ok=True)
        self.template_path.mkdir(parents=True, exist_ok=True)
        self.temp_path.mkdir(parents=True, exist_ok=True)

        # 任务管理
        self.active_tasks: Dict[str, ExportTask] = {}
        self._lock = threading.RLock()

        # 并发控制
        self.max_workers = self.config.get("max_workers", 4)
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)

        # 注册中文字体(PDF用)
        self._register_chinese_fonts()

        self.logger.info("增强导出服务初始化完成")

    def _register_chinese_fonts(self):
        """注册中文字体用于PDF"""
        try:
            if not REPORTLAB_AVAILABLE:
                return

            # 尝试注册常见的中文字体
            font_paths = [
                "C:/Windows/Fonts/simhei.ttf",  # Windows 黑体
                "C:/Windows/Fonts/simsun.ttc",  # Windows 宋体
                "/System/Library/Fonts/PingFang.ttc",  # macOS
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",  # Linux
            ]

            for font_path in font_paths:
                if Path(font_path).exists():
                    try:
                        pdfmetrics.registerFont(TTFont('CustomChinese', font_path))
                        self.logger.info(f"注册中文字体成功: {font_path}")
                        break
                    except Exception as e:
                        self.logger.warning(f"注册字体 {font_path} 失败: {e}")

        except Exception as e:
            self.logger.error(f"注册中文字体失败: {e}")

    async def create_export_task(
        self,
        task_name: str,
        export_format: str,
        template_type: str,
        data_source: Dict[str, Any],
        options: Dict[str, Any] = None
    ) -> Dict[str, Any]:
        """创建导出任务"""
        try:
            task_id = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
            options = options or {}

            # 生成输出路径
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            ext = self._get_file_extension(export_format)
            output_path = str(self.export_path / f"{task_name}_{timestamp}{ext}")

            # 创建任务
            task = ExportTask(
                task_id=task_id,
                task_name=task_name,
                export_format=export_format,
                template_type=template_type,
                data_source=data_source,
                output_path=output_path,
                created_at=datetime.now()
            )

            with self._lock:
                self.active_tasks[task_id] = task

            # 异步处理任务
            asyncio.create_task(self._process_export_task(task, options))

            return {
                "status": "success",
                "task_id": task_id,
                "message": f"导出任务创建成功: {task_name}"
            }

        except Exception as e:
            self.logger.error(f"创建导出任务失败: {e}")
            return {"status": "error", "error": str(e)}

    def _get_file_extension(self, export_format: str) -> str:
        """获取文件扩展名"""
        extensions = {
            ExportFormat.PDF: ".pdf",
            ExportFormat.WORD: ".docx",
            ExportFormat.EXCEL: ".xlsx",
            ExportFormat.GRAPH: ".png",
            ExportFormat.HTML: ".html",
            ExportFormat.JSON: ".json"
        }
        return extensions.get(export_format, ".txt")

    async def _process_export_task(self, task: ExportTask, options: Dict[str, Any]):
        """处理导出任务"""
        try:
            task.status = "running"
            task.progress = 10.0

            # 根据格式选择处理方法
            if task.export_format == ExportFormat.PDF:
                result = await self._export_audit_pdf(task, options)
            elif task.export_format == ExportFormat.WORD:
                result = await self._export_audit_word(task, options)
            elif task.export_format == ExportFormat.GRAPH:
                result = await self._export_relationship_graph(task, options)
            elif task.export_format == ExportFormat.EXCEL:
                result = await self._export_excel(task, options)
            else:
                raise ValueError(f"不支持的导出格式: {task.export_format}")

            if result["status"] == "success":
                task.status = "completed"
                task.progress = 100.0
                task.file_path = result["file_path"]
                task.file_size = result["file_size"]
            else:
                task.status = "failed"
                task.error_message = result.get("error", "未知错误")

            task.completed_at = datetime.now()

        except Exception as e:
            self.logger.error(f"处理导出任务失败: {e}")
            task.status = "failed"
            task.error_message = str(e)
            task.completed_at = datetime.now()

    async def _export_audit_pdf(self, task: ExportTask, options: Dict[str, Any]) -> Dict[str, Any]:
        """导出专业审计PDF报告"""
        try:
            if not REPORTLAB_AVAILABLE:
                raise ImportError("reportlab不可用，无法导出PDF")

            # 创建PDF文档
            doc = SimpleDocTemplate(
                task.output_path,
                pagesize=A4,
                topMargin=2*cm,
                bottomMargin=2*cm,
                leftMargin=2.5*cm,
                rightMargin=2.5*cm
            )

            story = []
            styles = self._create_pdf_styles()

            # 1. 封面页
            story.extend(self._create_pdf_cover_page(task, styles))
            story.append(PageBreak())

            # 2. 目录(可选)
            if options.get("include_toc", True):
                story.extend(self._create_pdf_toc(task, styles))
                story.append(PageBreak())

            # 3. 审计概要
            if task.template_type in [AuditReportTemplate.STANDARD, AuditReportTemplate.SUMMARY]:
                story.extend(self._create_audit_summary_section(task, styles))
                story.append(Spacer(1, 0.5*cm))

            # 4. 审计发现
            if "findings" in task.data_source:
                story.extend(self._create_findings_section(task, styles))
                story.append(Spacer(1, 0.5*cm))

            # 5. 证据详情
            if "evidence" in task.data_source and task.template_type == AuditReportTemplate.DETAILED:
                story.extend(self._create_evidence_section(task, styles))
                story.append(Spacer(1, 0.5*cm))

            # 6. 数据表格
            if "tables" in task.data_source:
                story.extend(self._create_data_tables_section(task, styles))
                story.append(Spacer(1, 0.5*cm))

            # 7. 图表分析
            if "charts" in task.data_source:
                story.extend(self._create_charts_section(task, styles))

            # 8. 结论与建议
            story.extend(self._create_conclusion_section(task, styles))

            # 生成PDF
            doc.build(story)

            file_size = Path(task.output_path).stat().st_size

            return {
                "status": "success",
                "file_path": task.output_path,
                "file_size": file_size,
                "format": "pdf"
            }

        except Exception as e:
            self.logger.error(f"导出PDF失败: {e}")
            return {"status": "error", "error": str(e)}

    def _create_pdf_styles(self) -> Dict[str, ParagraphStyle]:
        """创建PDF样式"""
        base_styles = getSampleStyleSheet()

        # 尝试使用中文字体
        try:
            font_name = 'CustomChinese'
            pdfmetrics.getFont(font_name)
        except:
            font_name = 'Helvetica'

        styles = {
            'Title': ParagraphStyle(
                'CustomTitle',
                parent=base_styles['Title'],
                fontName=font_name,
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                alignment=TA_CENTER,
                spaceAfter=30
            ),
            'Heading1': ParagraphStyle(
                'CustomH1',
                parent=base_styles['Heading1'],
                fontName=font_name,
                fontSize=18,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=12
            ),
            'Heading2': ParagraphStyle(
                'CustomH2',
                parent=base_styles['Heading2'],
                fontName=font_name,
                fontSize=14,
                textColor=colors.HexColor('#34495e'),
                spaceAfter=10,
                spaceBefore=10
            ),
            'Normal': ParagraphStyle(
                'CustomNormal',
                parent=base_styles['Normal'],
                fontName=font_name,
                fontSize=10,
                leading=14,
                textColor=colors.HexColor('#333333')
            ),
            'TableHeader': ParagraphStyle(
                'TableHeader',
                parent=base_styles['Normal'],
                fontName=font_name,
                fontSize=10,
                textColor=colors.white,
                alignment=TA_CENTER
            ),
            'TableCell': ParagraphStyle(
                'TableCell',
                parent=base_styles['Normal'],
                fontName=font_name,
                fontSize=9,
                leading=12
            )
        }

        return styles

    def _create_pdf_cover_page(self, task: ExportTask, styles: Dict) -> List:
        """创建PDF封面页"""
        elements = []

        # 标题
        title_text = task.data_source.get("report_title", task.task_name)
        elements.append(Spacer(1, 3*cm))
        elements.append(Paragraph(title_text, styles['Title']))
        elements.append(Spacer(1, 2*cm))

        # 基本信息
        info_data = [
            ['审计项目:', task.data_source.get("project_name", "N/A")],
            ['审计期间:', task.data_source.get("audit_period", "N/A")],
            ['审计人员:', task.data_source.get("auditor", "N/A")],
            ['报告日期:', datetime.now().strftime('%Y年%m月%d日')],
        ]

        info_table = Table(info_data, colWidths=[4*cm, 10*cm])
        info_table.setStyle(TableStyle([
            ('FONTNAME', (0, 0), (-1, -1), 'CustomChinese'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#555555')),
            ('TEXTCOLOR', (1, 0), (1, -1), colors.HexColor('#000000')),
            ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
            ('ALIGN', (1, 0), (1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ]))

        elements.append(info_table)

        return elements

    def _create_pdf_toc(self, task: ExportTask, styles: Dict) -> List:
        """创建目录"""
        elements = []

        elements.append(Paragraph("目录", styles['Heading1']))
        elements.append(Spacer(1, 0.5*cm))

        toc_items = [
            "1. 审计概要",
            "2. 审计发现",
            "3. 证据详情",
            "4. 数据分析",
            "5. 结论与建议"
        ]

        for item in toc_items:
            elements.append(Paragraph(item, styles['Normal']))
            elements.append(Spacer(1, 0.3*cm))

        return elements

    def _create_audit_summary_section(self, task: ExportTask, styles: Dict) -> List:
        """创建审计概要章节"""
        elements = []

        elements.append(Paragraph("一、审计概要", styles['Heading1']))
        elements.append(Spacer(1, 0.3*cm))

        summary = task.data_source.get("summary", {})

        # 审计目标
        elements.append(Paragraph("1.1 审计目标", styles['Heading2']))
        objective_text = summary.get("objective", "本次审计旨在评估财务数据的准确性和完整性。")
        elements.append(Paragraph(objective_text, styles['Normal']))
        elements.append(Spacer(1, 0.3*cm))

        # 审计范围
        elements.append(Paragraph("1.2 审计范围", styles['Heading2']))
        scope_text = summary.get("scope", "审计范围涵盖2023年度全部财务数据。")
        elements.append(Paragraph(scope_text, styles['Normal']))
        elements.append(Spacer(1, 0.3*cm))

        # 关键统计
        if "statistics" in summary:
            elements.append(Paragraph("1.3 关键统计", styles['Heading2']))
            stats = summary["statistics"]

            stats_data = [
                ['指标', '数值'],
                ['审计记录总数', str(stats.get("total_records", 0))],
                ['发现问题数', str(stats.get("issues_found", 0))],
                ['高风险问题', str(stats.get("high_risk", 0))],
                ['审计覆盖率', f"{stats.get('coverage', 0)}%"]
            ]

            stats_table = Table(stats_data, colWidths=[6*cm, 6*cm])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, -1), 'CustomChinese'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f2f2f2')])
            ]))

            elements.append(stats_table)

        return elements

    def _create_findings_section(self, task: ExportTask, styles: Dict) -> List:
        """创建审计发现章节"""
        elements = []

        elements.append(Paragraph("二、审计发现", styles['Heading1']))
        elements.append(Spacer(1, 0.3*cm))

        findings = task.data_source.get("findings", [])

        for idx, finding in enumerate(findings, 1):
            # 发现标题
            title = f"{idx}. {finding.get('title', '审计发现')}"
            elements.append(Paragraph(title, styles['Heading2']))

            # 发现描述
            description = finding.get("description", "")
            elements.append(Paragraph(f"<b>描述：</b>{description}", styles['Normal']))
            elements.append(Spacer(1, 0.2*cm))

            # 风险级别
            risk_level = finding.get("risk_level", "中")
            risk_colors = {"高": "#e74c3c", "中": "#f39c12", "低": "#27ae60"}
            risk_color = risk_colors.get(risk_level, "#95a5a6")
            elements.append(Paragraph(
                f"<b>风险级别：</b><font color='{risk_color}'>{risk_level}</font>",
                styles['Normal']
            ))
            elements.append(Spacer(1, 0.2*cm))

            # 建议措施
            recommendation = finding.get("recommendation", "")
            elements.append(Paragraph(f"<b>建议措施：</b>{recommendation}", styles['Normal']))
            elements.append(Spacer(1, 0.5*cm))

        return elements

    def _create_evidence_section(self, task: ExportTask, styles: Dict) -> List:
        """创建证据详情章节"""
        elements = []

        elements.append(Paragraph("三、证据详情", styles['Heading1']))
        elements.append(Spacer(1, 0.3*cm))

        evidence_list = task.data_source.get("evidence", [])

        for idx, evidence in enumerate(evidence_list, 1):
            elements.append(Paragraph(f"证据 #{idx}", styles['Heading2']))

            evidence_data = [
                ['项目', '内容'],
                ['证据编号', evidence.get('evidence_id', 'N/A')],
                ['证据类型', evidence.get('evidence_type', 'N/A')],
                ['收集日期', evidence.get('collection_date', 'N/A')],
                ['来源', evidence.get('source', 'N/A')],
                ['描述', evidence.get('description', 'N/A')[:100]]
            ]

            evidence_table = Table(evidence_data, colWidths=[4*cm, 10*cm])
            evidence_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, -1), 'CustomChinese'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#ecf0f1')])
            ]))

            elements.append(evidence_table)
            elements.append(Spacer(1, 0.4*cm))

        return elements

    def _create_data_tables_section(self, task: ExportTask, styles: Dict) -> List:
        """创建数据表格章节"""
        elements = []

        elements.append(Paragraph("四、数据分析", styles['Heading1']))
        elements.append(Spacer(1, 0.3*cm))

        tables = task.data_source.get("tables", {})

        for table_name, table_df in tables.items():
            if isinstance(table_df, pd.DataFrame) and not table_df.empty:
                elements.append(Paragraph(table_name, styles['Heading2']))

                # 限制行数
                display_df = table_df.head(20)

                # 构建表格数据
                table_data = [display_df.columns.tolist()]
                for _, row in display_df.iterrows():
                    table_data.append([str(val)[:30] for val in row])  # 限制单元格长度

                # 动态调整列宽
                col_count = len(display_df.columns)
                col_width = min(14 / col_count, 4) * cm

                pdf_table = Table(table_data, colWidths=[col_width] * col_count)
                pdf_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('FONTNAME', (0, 0), (-1, -1), 'CustomChinese'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')])
                ]))

                elements.append(pdf_table)
                elements.append(Spacer(1, 0.5*cm))

        return elements

    def _create_charts_section(self, task: ExportTask, styles: Dict) -> List:
        """创建图表章节"""
        elements = []

        # 如果有图表图片路径，添加到PDF
        charts = task.data_source.get("charts", [])

        if charts:
            elements.append(Paragraph("五、图表分析", styles['Heading1']))
            elements.append(Spacer(1, 0.3*cm))

            for chart in charts:
                if "image_path" in chart and Path(chart["image_path"]).exists():
                    elements.append(Paragraph(chart.get("title", "图表"), styles['Heading2']))

                    img = Image(chart["image_path"], width=14*cm, height=10*cm)
                    elements.append(img)
                    elements.append(Spacer(1, 0.5*cm))

        return elements

    def _create_conclusion_section(self, task: ExportTask, styles: Dict) -> List:
        """创建结论章节"""
        elements = []

        elements.append(Paragraph("五、结论与建议", styles['Heading1']))
        elements.append(Spacer(1, 0.3*cm))

        conclusion = task.data_source.get("conclusion", {})

        # 总体结论
        conclusion_text = conclusion.get(
            "summary",
            "基于本次审计工作，我们认为被审计单位的财务数据总体准确可靠。"
        )
        elements.append(Paragraph(conclusion_text, styles['Normal']))
        elements.append(Spacer(1, 0.5*cm))

        # 主要建议
        recommendations = conclusion.get("recommendations", [])
        if recommendations:
            elements.append(Paragraph("主要建议：", styles['Heading2']))
            for idx, rec in enumerate(recommendations, 1):
                elements.append(Paragraph(f"{idx}. {rec}", styles['Normal']))
                elements.append(Spacer(1, 0.2*cm))

        return elements

    async def _export_audit_word(self, task: ExportTask, options: Dict[str, Any]) -> Dict[str, Any]:
        """导出专业审计Word文档"""
        try:
            if not PYTHON_DOCX_AVAILABLE:
                raise ImportError("python-docx不可用，无法导出Word")

            # 创建文档
            doc = Document()

            # 设置页面属性
            section = doc.sections[0]
            section.page_height = Cm(29.7)  # A4高度
            section.page_width = Cm(21)     # A4宽度
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.18)
            section.right_margin = Cm(3.18)

            # 1. 标题
            title = doc.add_heading(task.data_source.get("report_title", task.task_name), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 2. 基本信息表
            doc.add_paragraph()
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Light Grid Accent 1'

            info_data = [
                ('审计项目', task.data_source.get("project_name", "N/A")),
                ('审计期间', task.data_source.get("audit_period", "N/A")),
                ('审计人员', task.data_source.get("auditor", "N/A")),
                ('报告日期', datetime.now().strftime('%Y年%m月%d日'))
            ]

            for idx, (label, value) in enumerate(info_data):
                info_table.rows[idx].cells[0].text = label
                info_table.rows[idx].cells[1].text = value

            doc.add_page_break()

            # 3. 审计概要
            doc.add_heading('一、审计概要', 1)
            summary = task.data_source.get("summary", {})
            doc.add_paragraph(summary.get("objective", "审计目标说明..."))

            # 4. 审计发现
            if "findings" in task.data_source:
                doc.add_heading('二、审计发现', 1)

                for idx, finding in enumerate(task.data_source["findings"], 1):
                    doc.add_heading(f'{idx}. {finding.get("title", "审计发现")}', 2)

                    # 发现详情表
                    finding_table = doc.add_table(rows=4, cols=2)
                    finding_table.style = 'Light List Accent 1'

                    finding_table.rows[0].cells[0].text = '描述'
                    finding_table.rows[0].cells[1].text = finding.get('description', '')
                    finding_table.rows[1].cells[0].text = '风险级别'
                    finding_table.rows[1].cells[1].text = finding.get('risk_level', '中')
                    finding_table.rows[2].cells[0].text = '影响金额'
                    finding_table.rows[2].cells[1].text = str(finding.get('amount', 'N/A'))
                    finding_table.rows[3].cells[0].text = '建议措施'
                    finding_table.rows[3].cells[1].text = finding.get('recommendation', '')

                    doc.add_paragraph()

            # 5. 数据表格
            if "tables" in task.data_source:
                doc.add_heading('三、数据分析', 1)

                for table_name, table_df in task.data_source["tables"].items():
                    if isinstance(table_df, pd.DataFrame) and not table_df.empty:
                        doc.add_heading(table_name, 2)

                        # 限制行数
                        display_df = table_df.head(50)

                        # 创建表格
                        word_table = doc.add_table(rows=len(display_df) + 1, cols=len(display_df.columns))
                        word_table.style = 'Light Grid Accent 1'

                        # 表头
                        for col_idx, col_name in enumerate(display_df.columns):
                            word_table.rows[0].cells[col_idx].text = str(col_name)

                        # 数据行
                        for row_idx, (_, row) in enumerate(display_df.iterrows(), 1):
                            for col_idx, value in enumerate(row):
                                word_table.rows[row_idx].cells[col_idx].text = str(value)[:100]

                        doc.add_paragraph()

            # 6. 结论
            doc.add_heading('四、结论与建议', 1)
            conclusion = task.data_source.get("conclusion", {})
            doc.add_paragraph(conclusion.get("summary", "审计结论..."))

            if "recommendations" in conclusion:
                doc.add_heading('主要建议', 2)
                for idx, rec in enumerate(conclusion["recommendations"], 1):
                    doc.add_paragraph(f'{idx}. {rec}', style='List Number')

            # 保存文档
            doc.save(task.output_path)

            file_size = Path(task.output_path).stat().st_size

            return {
                "status": "success",
                "file_path": task.output_path,
                "file_size": file_size,
                "format": "word"
            }

        except Exception as e:
            self.logger.error(f"导出Word失败: {e}")
            return {"status": "error", "error": str(e)}

    async def _export_relationship_graph(self, task: ExportTask, options: Dict[str, Any]) -> Dict[str, Any]:
        """导出证据关系图谱"""
        try:
            if not NETWORKX_AVAILABLE:
                raise ImportError("networkx不可用，无法生成图谱")

            # 获取关系数据
            relationships = task.data_source.get("relationships", [])
            nodes = task.data_source.get("nodes", [])

            if not relationships or not nodes:
                raise ValueError("缺少图谱数据")

            # 创建图
            G = nx.DiGraph()

            # 添加节点
            for node in nodes:
                G.add_node(
                    node["id"],
                    label=node.get("label", node["id"]),
                    node_type=node.get("type", "default"),
                    size=node.get("size", 1000)
                )

            # 添加边
            for rel in relationships:
                G.add_edge(
                    rel["source"],
                    rel["target"],
                    label=rel.get("label", ""),
                    weight=rel.get("weight", 1.0)
                )

            # 绘制图谱
            plt.figure(figsize=(16, 12))
            plt.clf()

            # 使用spring布局
            pos = nx.spring_layout(G, k=2, iterations=50, seed=42)

            # 节点颜色映射
            node_types = nx.get_node_attributes(G, 'node_type')
            color_map = {
                'evidence': '#3498db',
                'finding': '#e74c3c',
                'transaction': '#2ecc71',
                'account': '#f39c12',
                'default': '#95a5a6'
            }
            node_colors = [color_map.get(node_types.get(node, 'default'), '#95a5a6') for node in G.nodes()]

            # 绘制节点
            nx.draw_networkx_nodes(
                G, pos,
                node_color=node_colors,
                node_size=3000,
                alpha=0.8,
                edgecolors='white',
                linewidths=2
            )

            # 绘制边
            nx.draw_networkx_edges(
                G, pos,
                edge_color='#7f8c8d',
                width=2,
                alpha=0.5,
                arrows=True,
                arrowsize=20,
                arrowstyle='->'
            )

            # 绘制标签
            labels = nx.get_node_attributes(G, 'label')
            nx.draw_networkx_labels(
                G, pos,
                labels,
                font_size=10,
                font_family='SimHei',
                font_color='white',
                font_weight='bold'
            )

            # 边标签
            edge_labels = nx.get_edge_attributes(G, 'label')
            if edge_labels:
                nx.draw_networkx_edge_labels(
                    G, pos,
                    edge_labels,
                    font_size=8,
                    font_family='SimHei'
                )

            plt.title(task.task_name, fontsize=18, fontweight='bold', pad=20)
            plt.axis('off')
            plt.tight_layout()

            # 保存图片
            plt.savefig(
                task.output_path,
                dpi=300,
                bbox_inches='tight',
                facecolor='white',
                edgecolor='none'
            )
            plt.close()

            file_size = Path(task.output_path).stat().st_size

            return {
                "status": "success",
                "file_path": task.output_path,
                "file_size": file_size,
                "format": "graph",
                "graph_stats": {
                    "nodes": len(G.nodes()),
                    "edges": len(G.edges()),
                    "density": nx.density(G)
                }
            }

        except Exception as e:
            self.logger.error(f"导出关系图谱失败: {e}")
            return {"status": "error", "error": str(e)}

    async def _export_excel(self, task: ExportTask, options: Dict[str, Any]) -> Dict[str, Any]:
        """导出Excel(调用原有功能)"""
        try:
            # 这里可以调用 multi_format_exporter 的Excel导出功能
            # 或者实现简单的Excel导出

            import openpyxl
            from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

            wb = openpyxl.Workbook()

            # 移除默认sheet
            wb.remove(wb.active)

            # 导出每个表
            tables = task.data_source.get("tables", {})
            for sheet_name, table_df in tables.items():
                if isinstance(table_df, pd.DataFrame) and not table_df.empty:
                    ws = wb.create_sheet(title=sheet_name[:31])  # Excel限制31字符

                    # 写入表头
                    for col_idx, col_name in enumerate(table_df.columns, 1):
                        cell = ws.cell(row=1, column=col_idx)
                        cell.value = col_name
                        cell.font = Font(bold=True, color="FFFFFF")
                        cell.fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
                        cell.alignment = Alignment(horizontal="center", vertical="center")

                    # 写入数据
                    for row_idx, (_, row) in enumerate(table_df.iterrows(), 2):
                        for col_idx, value in enumerate(row, 1):
                            ws.cell(row=row_idx, column=col_idx, value=value)

                    # 自动调整列宽
                    for column in ws.columns:
                        max_length = 0
                        column_letter = column[0].column_letter
                        for cell in column:
                            try:
                                if len(str(cell.value)) > max_length:
                                    max_length = len(str(cell.value))
                            except:
                                pass
                        adjusted_width = min(max_length + 2, 50)
                        ws.column_dimensions[column_letter].width = adjusted_width

                    # 冻结首行
                    ws.freeze_panes = "A2"

            # 保存
            wb.save(task.output_path)

            file_size = Path(task.output_path).stat().st_size

            return {
                "status": "success",
                "file_path": task.output_path,
                "file_size": file_size,
                "format": "excel"
            }

        except Exception as e:
            self.logger.error(f"导出Excel失败: {e}")
            return {"status": "error", "error": str(e)}

    async def batch_export(
        self,
        export_configs: List[Dict[str, Any]]
    ) -> Dict[str, Any]:
        """批量导出"""
        try:
            task_ids = []

            for config in export_configs:
                result = await self.create_export_task(
                    task_name=config["task_name"],
                    export_format=config["export_format"],
                    template_type=config.get("template_type", AuditReportTemplate.STANDARD),
                    data_source=config["data_source"],
                    options=config.get("options", {})
                )

                if result["status"] == "success":
                    task_ids.append(result["task_id"])

            return {
                "status": "success",
                "message": f"批量导出任务创建成功，共 {len(task_ids)} 个任务",
                "task_ids": task_ids
            }

        except Exception as e:
            self.logger.error(f"批量导出失败: {e}")
            return {"status": "error", "error": str(e)}

    async def get_task_status(self, task_id: str) -> Dict[str, Any]:
        """获取任务状态"""
        try:
            with self._lock:
                if task_id not in self.active_tasks:
                    return {"status": "error", "error": f"任务 {task_id} 不存在"}

                task = self.active_tasks[task_id]

                return {
                    "status": "success",
                    "task": {
                        "task_id": task.task_id,
                        "task_name": task.task_name,
                        "export_format": task.export_format,
                        "status": task.status,
                        "progress": task.progress,
                        "file_path": task.file_path,
                        "file_size": task.file_size,
                        "error_message": task.error_message,
                        "created_at": task.created_at.isoformat() if task.created_at else None,
                        "completed_at": task.completed_at.isoformat() if task.completed_at else None
                    }
                }

        except Exception as e:
            self.logger.error(f"获取任务状态失败: {e}")
            return {"status": "error", "error": str(e)}

    async def cleanup(self):
        """清理资源"""
        try:
            if hasattr(self, 'executor'):
                self.executor.shutdown(wait=True)

            self.logger.info("增强导出服务清理完成")

        except Exception as e:
            self.logger.error(f"清理资源失败: {e}")


# 测试用例
async def test_enhanced_export():
    """测试增强导出功能"""

    service = EnhancedExportService({
        "export_path": "test_exports/",
        "template_path": "templates/",
        "temp_path": "temp/"
    })

    # 测试数据
    test_data = {
        "report_title": "2023年度财务审计报告",
        "project_name": "XX公司年度审计",
        "audit_period": "2023年1月-12月",
        "auditor": "审计团队",
        "summary": {
            "objective": "本次审计旨在评估XX公司2023年度财务报表的真实性、准确性和完整性。",
            "scope": "审计范围涵盖公司全部财务数据，包括资产负债表、利润表和现金流量表。",
            "statistics": {
                "total_records": 15234,
                "issues_found": 12,
                "high_risk": 3,
                "coverage": 98.5
            }
        },
        "findings": [
            {
                "title": "应收账款账龄异常",
                "description": "发现部分应收账款账龄超过2年仍未收回，存在坏账风险。",
                "risk_level": "高",
                "amount": 850000,
                "recommendation": "建议及时清理长期应收账款，计提坏账准备。"
            },
            {
                "title": "费用报销流程不完善",
                "description": "部分费用报销缺少完整的审批流程和原始凭证。",
                "risk_level": "中",
                "amount": 120000,
                "recommendation": "建议完善费用报销制度，加强内部控制。"
            }
        ],
        "tables": {
            "资产负债表": pd.DataFrame({
                "科目": ["货币资金", "应收账款", "存货", "固定资产"],
                "期末余额": [5000000, 2000000, 1500000, 10000000],
                "期初余额": [4500000, 2200000, 1400000, 10500000],
                "变动额": [500000, -200000, 100000, -500000]
            }),
            "利润表": pd.DataFrame({
                "项目": ["营业收入", "营业成本", "管理费用", "净利润"],
                "本期金额": [50000000, 35000000, 5000000, 8000000],
                "上期金额": [45000000, 32000000, 4500000, 7000000]
            })
        },
        "conclusion": {
            "summary": "基于本次审计工作，我们认为XX公司2023年度财务报表在所有重大方面公允反映了公司的财务状况和经营成果。",
            "recommendations": [
                "加强应收账款管理，建立催收机制",
                "完善内部控制制度，特别是费用报销流程",
                "定期进行资产盘点，确保账实相符"
            ]
        }
    }

    # 测试PDF导出
    print("测试PDF导出...")
    pdf_result = await service.create_export_task(
        task_name="财务审计报告",
        export_format=ExportFormat.PDF,
        template_type=AuditReportTemplate.DETAILED,
        data_source=test_data
    )
    print(f"PDF导出: {json.dumps(pdf_result, indent=2, ensure_ascii=False)}")

    # 等待完成
    await asyncio.sleep(5)

    # 测试Word导出
    print("\n测试Word导出...")
    word_result = await service.create_export_task(
        task_name="财务审计报告",
        export_format=ExportFormat.WORD,
        template_type=AuditReportTemplate.STANDARD,
        data_source=test_data
    )
    print(f"Word导出: {json.dumps(word_result, indent=2, ensure_ascii=False)}")

    # 等待完成
    await asyncio.sleep(5)

    # 测试关系图谱导出
    print("\n测试关系图谱导出...")
    graph_data = {
        "report_title": "审计证据关系图谱",
        "nodes": [
            {"id": "E001", "label": "证据1", "type": "evidence"},
            {"id": "E002", "label": "证据2", "type": "evidence"},
            {"id": "F001", "label": "发现1", "type": "finding"},
            {"id": "T001", "label": "交易1", "type": "transaction"},
        ],
        "relationships": [
            {"source": "E001", "target": "F001", "label": "支持"},
            {"source": "E002", "target": "F001", "label": "支持"},
            {"source": "T001", "target": "E001", "label": "来源"},
        ]
    }

    graph_result = await service.create_export_task(
        task_name="证据关系图谱",
        export_format=ExportFormat.GRAPH,
        template_type=AuditReportTemplate.EVIDENCE,
        data_source=graph_data
    )
    print(f"图谱导出: {json.dumps(graph_result, indent=2, ensure_ascii=False)}")

    await asyncio.sleep(5)

    await service.cleanup()

    print("\n✅ 所有测试完成！")


if __name__ == "__main__":
    asyncio.run(test_enhanced_export())
